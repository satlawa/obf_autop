import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib as mpl
import docx

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.shared import Pt

from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION

from OBFFlaechentabelle import OBFFlaechentabelle
from OBFDictionary import OBFDictionary
from OBFDictionary import OBFDictionary
from OBFDocX import OBFDocX

class OBFFlTab(object):

    def __init__(self, data_path, to_now):

        self.data_path = data_path
        self.to_now = to_now
        self.check_files_print()

        self.run()


    def check_files_print(self):
        '''
            check and print out if data files exist
        '''
        print('-----------------------------------------------------')
        print('-                                                   -')
        print('-                   files exist                     -')
        print('-                                                   -')
        print('-----------------------------------------------------')
        print('SAP Rohdaten:       to_' + self.to_now + '_sap.XLS:             ' + str(os.path.isfile(os.path.join(self.data_path, 'TO' + self.to_now, 'to_' + self.to_now + '_sap.XLS'))))
        print('SAP Flurnamen:      to_' + self.to_now + '_flurnamen.xlsx:      ' + str(os.path.isfile(os.path.join(self.data_path, 'TO' + self.to_now, 'to_' + self.to_now + '_flurnamen.xlsx'))))
        print('-----------------------------------------------------')
        print('Zusatz Dateien:     dict:                        ' + str(self.check_files_background()))
        print('-----------------------------------------------------')


    def check_files_background(self):
        '''
            check if necessary background files exist
            out: boolean
        '''
        exists_1 = os.path.isfile(os.path.join(self.data_path, 'dict', '3_data_Standortseinheiten.xlsx'))
        exists_2 = os.path.isfile(os.path.join(self.data_path, 'dict', '3_data_Wuchsgebiete.xlsx'))
        exists_3 = os.path.isfile(os.path.join(self.data_path, 'dict', '3_data_Wuchsgebiete_hoehenstufen.xlsx'))
        exists_4 = os.path.isfile(os.path.join(self.data_path, 'dict', '3_wuchsgebiete.png'))
        exists_5 = os.path.isfile(os.path.join(self.data_path, 'dict', 'auswertekat_index.xlsx'))
        exists_6 = os.path.isfile(os.path.join(self.data_path, 'dict', 'Ertragstafelsets.xlsx'))
        exists_7 = os.path.isfile(os.path.join(self.data_path, 'dict', 'key.xlsx'))
        exists_8 = os.path.isfile(os.path.join(self.data_path, 'dict', 'oebf.png'))
        exists_9 = os.path.isfile(os.path.join(self.data_path, 'dict', 'templet_xx.docx'))
        exists = exists_1 & exists_2 & exists_3 & exists_4 & exists_5 & exists_6 & exists_7 & exists_8 & exists_9

        return exists


    def run(self):
        '''
            run program
        '''

        version = '3.8'

        #-----------------------------------------------------------------------------------------
        #--- check if data exists
        #-----------------------------------------------------------------------------------------
        # background data
        if self.check_files_background() == False:
            return('background data is missing, the TO could not be made')
        # Taxationsdaten
        dat_fuc = os.path.isfile(os.path.join(self.data_path, 'TO' + self.to_now, 'to_' + self.to_now + '_sap.XLS'))                # SAP roh

        #-----------------------------------------------------------------------------------------
        #--- get data
        #-----------------------------------------------------------------------------------------
        path_dir = os.path.join(self.data_path, 'TO' + self.to_now)
        path_dict = os.path.join(self.data_path, 'dict')
        path_flur = os.path.join(self.data_path, 'flurnamen')

        #path_data = path_dir + '/to_' + self.to_now + '_sap.XLS'
        path_data = path_dir + '/to_' + self.to_now + '_sap.XLS'

        obf_flt = OBFFlaechentabelle(pd.read_csv(path_data, sep='\t', encoding = "ISO-8859-1", decimal=',', error_bad_lines=False))

        obf_dic = OBFDictionary()

        obf_doc = OBFDocX(Document(path_dict + '/templet_xx.docx'))

#xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx
        path_flur = os.path.join(path_dir, 'to_' + self.to_now + '_flurnamen.xlsx')
        data_flur = pd.read_excel(path_flur)
        data_flur = data_flur.fillna('')


        #######################
        ###   0 Deckblatt   ###
        #######################
        print('***   0 Deckblatt   ***')

        obf_doc.doc.add_picture(path_dict + '/oebf.png')

        for i in range(4):
            obf_doc.doc.add_paragraph('')

        p = obf_doc.doc.add_paragraph().add_run('Flächentabelle zum OPERAT')
        p.bold = True
        p.font.name = 'Verdana'
        p.font.size = Pt(34)

        p = obf_doc.doc.add_paragraph().add_run('FB ' + str(obf_flt.dic.fb) + ' ' + obf_flt.dic.dic_num_fb[obf_flt.dic.fb])
        p.bold = True
        p.font.name = 'Verdana'
        p.font.size = Pt(35)

        obf_doc.doc.add_paragraph('')

        p = obf_doc.doc.add_paragraph().add_run('Teiloperat ' + str(obf_flt.dic.to)) #+ str(obf_flt.dic.to))
        p.font.name = 'Verdana'
        p.font.size = Pt(25)

        for i in range(3):
            obf_doc.doc.add_paragraph('')

        for fr in obf_flt.dic.fr:

            p = obf_doc.doc.add_paragraph().add_run('FR ' + str(fr) + ' ' + obf_flt.dic.dic_num_fr[fr])
            p.font.name = 'Verdana'
            p.font.size = Pt(20)

        for i in range(12-(obf_flt.dic.fr.shape[0]*2)):
            p = obf_doc.doc.add_paragraph().add_run('')
            p.font.size = Pt(20)

        p = obf_doc.doc.add_paragraph().add_run('Organisationsstand 01.01.' + obf_flt.dic.laufzeit_begin)
        p.font.name = 'Verdana'
        p.font.size = Pt(18)

        p = obf_doc.doc.add_paragraph().add_run('Laufzeit 01.01.' + obf_flt.dic.laufzeit_begin + ' - 31.12.' + obf_flt.dic.laufzeit_end)
        p.font.name = 'Verdana'
        p.font.size = Pt(18)

        for i in range(10-(obf_flt.dic.fr.shape[0]*2)):
            obf_doc.doc.add_paragraph('')

        p = obf_doc.doc.add_paragraph().add_run('Für die Unternehmensleitung                       Für den Forstbetrieb')
        p.font.name = 'Verdana'
        p.font.size = Pt(12)

        obf_doc.doc.add_page_break()

        #############################
        ###   1 Abkürzungsverzeichnis   ###
        #############################

        for i in range(3):
            obf_doc.doc.add_paragraph('')

        run = obf_doc.doc.add_paragraph('Erstellung der Flächentabellen mit AutOP ').add_run(' Version ' + version)
        font = run.font
        font.name = 'Verdana'
        font.size = Pt(8)
        obf_doc.doc.add_paragraph('')
        obf_doc.doc.add_paragraph('')

        print('***   1 Abkürzungsverzeichnis   ***')
        obf_doc.doc.add_heading('Abkürzungsverzeichnis', 1)

        obf_doc.doc.add_paragraph('')
        obf_doc.doc.add_paragraph('Die folgenden Abkürzungen werden in den Tabellen verwendet, um die Lesbarkeit und Übersicht der Tabellen zu gewährleisten.')
        obf_doc.doc.add_paragraph('')

        obf_doc.doc.add_heading('Spaltenüberschriften', 2)

        obf_doc.doc.add_paragraph('')
        obf_doc.doc.add_paragraph('Alle nachfolgenden Tabellen verwenden diese Abkürzungen in den Spaltenüberschriften.')
        obf_doc.doc.add_paragraph('')
        obf_doc.doc.add_paragraph('Abt		...	Abteilung')
        obf_doc.doc.add_paragraph('UAbt		...	Unterabteilung')
        obf_doc.doc.add_paragraph('TF		...	Teilfläche')
        obf_doc.doc.add_paragraph('WE		...	Wirtschaftseinheits-Typ')
        obf_doc.doc.add_paragraph('BKL		...	Betriebsklasse')
        obf_doc.doc.add_paragraph('UZ		...	Umtriebszeit')
        obf_doc.doc.add_paragraph('ES		...	Ertragssituation')
        obf_doc.doc.add_paragraph('BW		...	Bewirtschaftungsform')
        obf_doc.doc.add_paragraph('SW		...	Schutzwaldkategorie')
        obf_doc.doc.add_paragraph('NG		...	Nebengrund Art')
        obf_doc.doc.add_paragraph('Fl WW		...	Wirtschaftswaldfläche')
        obf_doc.doc.add_paragraph('Fl SW		...	Schutzwaldfläche')
        obf_doc.doc.add_paragraph('Fl NHB		...	Nichtholzbodenfläche')
        obf_doc.doc.add_paragraph('Fl pNG		...	produktiver Nebengrund')
        obf_doc.doc.add_paragraph('Fl uNG		...	unproduktiver Nebengrund')
        obf_doc.doc.add_paragraph('Fl Ges		...	Gesamtfläche')
        obf_doc.doc.add_paragraph('')

        obf_doc.doc.add_page_break()

        obf_doc.doc.add_heading('Abkürzungen von Werten', 2)

        obf_doc.doc.add_paragraph('')
        obf_doc.doc.add_paragraph('Alle nachfolgenden Tabellen verwenden diese Abkürzungen von Werten innerhalb der Tabellen.')
        obf_doc.doc.add_paragraph('')

        p = obf_doc.doc.add_paragraph().add_run('Wirtschaftseinheits-Typ')
        p.font.bold = True
        obf_doc.doc.add_paragraph('')
        obf_doc.doc.add_paragraph('WO	...	Waldort')
        obf_doc.doc.add_paragraph('NG	...	Nebengrund')
        obf_doc.doc.add_paragraph('')

        p = obf_doc.doc.add_paragraph().add_run('Ertragssituation')
        p.font.bold = True
        obf_doc.doc.add_paragraph('')
        obf_doc.doc.add_paragraph('I	...	in Ertrag')
        obf_doc.doc.add_paragraph('A	...	außer Ertrag')
        obf_doc.doc.add_paragraph('')

        p = obf_doc.doc.add_paragraph().add_run('Bewirtschaftungsform')
        p.font.bold = True
        obf_doc.doc.add_paragraph('')
        obf_doc.doc.add_paragraph('W	...	Wirtschaftswald')
        obf_doc.doc.add_paragraph('S	...	Schutzwald')
        obf_doc.doc.add_paragraph('')

        p = obf_doc.doc.add_paragraph().add_run('Schutzwaldkategorie')
        p.font.bold = True
        obf_doc.doc.add_paragraph('')
        obf_doc.doc.add_paragraph('S	...	Standortsschutzwald')
        obf_doc.doc.add_paragraph('O	...	Objektschutzwald')
        obf_doc.doc.add_paragraph('B	...	Bannwald')
        obf_doc.doc.add_paragraph('')

        p = obf_doc.doc.add_paragraph().add_run('Nebengrund Art')
        p.font.bold = True
        obf_doc.doc.add_paragraph('')
        obf_doc.doc.add_paragraph('1	...	Feuchtbiotop')
        obf_doc.doc.add_paragraph('2	...	Latschenfelder')
        obf_doc.doc.add_paragraph('3	...	Forststraße')
        obf_doc.doc.add_paragraph('4	...	Sonstiger Nichtholzboden')
        obf_doc.doc.add_paragraph('5	...	Wiese')
        obf_doc.doc.add_paragraph('6	...	Acker')
        obf_doc.doc.add_paragraph('7	...	Gewässer')
        obf_doc.doc.add_paragraph('8	...	Ödflächen')
        obf_doc.doc.add_paragraph('9	...	Sonstiger unproduktiver Nebengrund')
        obf_doc.doc.add_paragraph('')

        p = obf_doc.doc.add_paragraph().add_run('Flächen')
        p.font.bold = True
        obf_doc.doc.add_paragraph('')
        obf_doc.doc.add_paragraph('Wirtschaftswaldfläche	...	Wirtschaftswald')
        obf_doc.doc.add_paragraph('Schutzwaldfläche		...	Schutzwald')
        obf_doc.doc.add_paragraph('Nichtholzbodenfläche		...	Nebengrund Art 3 & 4')
        obf_doc.doc.add_paragraph('produktiver Nebengrund	...	Nebengrund Art 5 & 6')
        obf_doc.doc.add_paragraph('produktiver Nebengrund	...	Nebengrund Art 1 & 2 & 7 & 8 & 9')
        obf_doc.doc.add_paragraph('Gesamtfläche			...	alle Flächen summiert')
        obf_doc.doc.add_paragraph('')


        obf_doc.doc.add_page_break()

        #############################
        ###   2 Flächentabellen   ###
        #############################

        print('***   2 Flächentabellen   ***')
        obf_doc.doc.add_heading('Flächentabellen', 1)
        obf_doc.doc.add_paragraph('')

        data_flur_fb = data_flur.loc[(data_flur['FB']==obf_flt.dic.fb)]

        for i,fr in enumerate(obf_flt.dic.fr):
            print('---------------------------   ' + str(fr) + '   ---------------------------')

            obf_doc.doc.add_heading('FR  ' + str(fr) + ' ' +obf_flt.dic.dic_num_fr[fr], 2)
            obf_doc.doc.add_paragraph('')

            data_flur_fr = data_flur_fb.loc[(data_flur_fb['FR']==fr)]

            for abt in obf_flt.fr_abt[i]:
                print(abt)
                table = obf_flt.create_table(fr, abt)

                # add Caption to table
                #obf_doc.docx_paragraph_table('Flächentabelle der Abteilung ' + str(abt) + ', im FR ' + str(fr) + ' ' + obf_flt.dic.dic_num_fr[fr] + ', im FB ' + str(obf_flt.dic.fb) + ' ' + obf_flt.dic.dic_num_fb[obf_flt.dic.fb] + ', Abkürzungen sind im Abkürzungsverzeichins aufgeschlüsselt.')

                # add table
                flur = data_flur_fr.loc[(data_flur_fr['ABT']==abt)].values[0,3]
                obf_doc.docx_table_x(table, width1=Cm(1), width2=Cm(1), header_rep=True, header = 'FR ' + str(fr) + '   |   Abteilung   ' + str(abt) + '   |   ' + flur + '   |   [ha]', font_size = 7, autofit = True)
                obf_doc.doc.add_paragraph('')
                obf_doc.doc.add_paragraph('')

            obf_doc.doc.add_page_break()

        # save the doc
        path_save = os.path.join(self.data_path, 'TO' + self.to_now + '_flaechentabelle_.docx')
        obf_doc.doc.save(path_save)

        print('end')
        print('successfully made in ' + self.data_path)
