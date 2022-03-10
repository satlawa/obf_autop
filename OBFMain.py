
import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib as mpl
import docx

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.shared import Pt

from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION


from OBFFunc import OBFFunc

from OBFHiebssatz import OBFHiebssatz
from OBFEinschlag import OBFEinschlag
from OBFEinschlagHiebsatz import OBFEinschlagHiebsatz
from OBFSchadholz import OBFSchadholz
from OBFWaldpflege import OBFWaldpflege
from OBFKlima import OBFKlima
from OBFNatur import OBFNatur
from OBFSpi import OBFSpi

from OBFEt import OBFEt
from OBFDocX import OBFDocX
from OBFText import OBFText
from OBFDictionary import OBFDictionary


class OBFMain(object):

    def __init__(self, data_path, to_now, to_old, klima_stationen, do_sections):

        self.data_path = data_path
        self.to_now = to_now
        self.to_old = to_old
        self.do_sections = do_sections
        self.klima_stationen = klima_stationen


    def save_log(self, log):
        '''
            save log
        '''
        with open("log.txt", "w") as text_file:
            print("{}".format(log), file=text_file)


    def check_files_print(self):
        '''
            check and print out if data files exist
        '''
        print('-----------------------------------------------------')
        print('-                                                   -')
        print('-                   files exist                     -')
        print('-                                                   -')
        print('-----------------------------------------------------')
        print('SAP Rohdaten:       to_' + self.to_now + '_sap.XLS:             ' + str(os.path.isfile(os.path.join(self.data_path, 'TO' + self.to_now, 'to_' + self.to_now + '_sap.XLS'))))
        print('SAP Ausertekat.:    to_' + self.to_now + '_sap_natur.XLS:       ' + str(os.path.isfile(os.path.join(self.data_path, 'TO' + self.to_now, 'to_' + self.to_now + '_sap_natur.XLS'))))
        print('-----------------------------------------------------')
        print('SAP HS-Bilanz neu:  to_' + self.to_now + '_hs_bilanz.XLS:       ' + str(os.path.isfile(os.path.join(self.data_path, 'TO' + self.to_now, 'to_' + self.to_now + '_hs_bilanz.XLS'))))
        print('SAP HS-Bilanz alt:  to_' + self.to_now + '_hs_bilanz_old.XLS:   ' + str(os.path.isfile(os.path.join(self.data_path, 'TO' + self.to_now, 'to_' + self.to_now + '_hs_bilanz_old.XLS'))))
        print('-----------------------------------------------------')
        print('BW Alter:           to_' + self.to_now + '_bw_alter.xlsx:       ' + str(os.path.isfile(os.path.join(self.data_path, 'TO' + self.to_now, 'to_' + self.to_now + '_bw_alter.xlsx'))))
        print('BW Neigung:         to_' + self.to_now + '_bw_neig.xlsx:        ' + str(os.path.isfile(os.path.join(self.data_path, 'TO' + self.to_now, 'to_' + self.to_now + '_bw_neig.xlsx'))))
        print('BW Seehöhe:         to_' + self.to_now + '_bw_seeh.xlsx:        ' + str(os.path.isfile(os.path.join(self.data_path, 'TO' + self.to_now, 'to_' + self.to_now + '_bw_seeh.xlsx'))))
        print('BW Umtriebszeit:    to_' + self.to_now + '_bw_uz.xlsx:          ' + str(os.path.isfile(os.path.join(self.data_path, 'TO' + self.to_now, 'to_' + self.to_now + '_bw_uz.xlsx'))))
        print('BW Zufällige:       to_' + self.to_now + '_bw_zufaellige.xlsx:  ' + str(os.path.isfile(os.path.join(self.data_path, 'TO' + self.to_now, 'to_' + self.to_now + '_bw_zufaellige.xlsx'))))
        print('BW Waldpflege:      to_' + self.to_now + '_bw_wp.xlsx:          ' + str(os.path.isfile(os.path.join(self.data_path, 'TO' + self.to_now, 'to_' + self.to_now + '_bw_wp.xlsx'))))
        print('-----------------------------------------------------')
        print('SAP Stichprobe:     SPI_' + self.fe_year + '.txt:                ' + str(os.path.isfile(os.path.join(self.data_path, 'stichprobe', 'SPI_' + self.fe_year + '.txt'))))
        print('-----------------------------------------------------')
        print('Klima Temperatur:   1_Lufttemperatur.txt:        ' + str(os.path.isfile(os.path.join(self.data_path, 'klima', '1_Lufttemperatur.txt'))))
        print('Klima Niederschlag: 2_Niederschlag.txt:          ' + str(os.path.isfile(os.path.join(self.data_path, 'klima', '2_Niederschlag.txt'))))
        print('Klima Schnee:       4_Schnee.txt:                ' + str(os.path.isfile(os.path.join(self.data_path, 'klima', '4_Schnee.txt'))))
        print('-----------------------------------------------------')
        print('Zusatz Dateien:     dict:                        ' + str(self.check_files_background()))
        print('-----------------------------------------------------')


    def check_files_background(self):
        '''
            check if necessary background files exist
            out: boolean
        '''
        exists_1 = os.path.isfile(os.path.join(self.data_path, 'dict', '3_data_Standortseinheiten.xlsx'))
        exists_2 = os.path.isfile(os.path.join(self.data_path, 'dict', '3_data_Wuchsgebiete.xlsx'))
        exists_3 = os.path.isfile(os.path.join(self.data_path, 'dict', '3_data_Wuchsgebiete_hoehenstufen.xlsx'))
        exists_4 = os.path.isfile(os.path.join(self.data_path, 'dict', '3_wuchsgebiete.png'))
        exists_5 = os.path.isfile(os.path.join(self.data_path, 'dict', 'auswertekat_index.xlsx'))
        exists_6 = os.path.isfile(os.path.join(self.data_path, 'dict', 'Ertragstafelsets.xlsx'))
        exists_7 = os.path.isfile(os.path.join(self.data_path, 'dict', 'key.xlsx'))
        exists_8 = os.path.isfile(os.path.join(self.data_path, 'dict', 'oebf.png'))
        exists_9 = os.path.isfile(os.path.join(self.data_path, 'dict', 'templet_xx.docx'))
        exists = exists_1 & exists_2 & exists_3 & exists_4 & exists_5 & exists_6 & exists_7 & exists_8 & exists_9

        return exists


    def run(self):
        '''
            run program
        '''

        version = '3.9'

        #-----------------------------------------------------------------------------------------
        #--- check if data exists
        #-----------------------------------------------------------------------------------------

        # set path dir
        path_dir = os.path.join(self.data_path, 'TO' + self.to_now)
        path_dict = os.path.join(self.data_path, 'dict')

        # background data
        if self.check_files_background() == False:
            return('dict data is missing, the TO could not be made')

        # Taxationsdaten
        # check if file exists
        dat_fuc = os.path.isfile(os.path.join(self.data_path, 'TO' + self.to_now, 'to_' + self.to_now + '_sap.XLS'))                # SAP roh
        # get data
        path_data = path_dir + '/to_' + self.to_now + '_sap.XLS'
        obf_fuc = OBFFunc(pd.read_csv(path_data, sep='\t', encoding = "ISO-8859-1", decimal=',', error_bad_lines=False))
        # set fe_year
        self.fe_year = str(int(obf_fuc.data.loc[0,'Beg. Laufzeit'][-4:])-1)

        self.check_files_print()

        dat_natur = os.path.isfile(os.path.join(self.data_path, 'TO' + self.to_now, 'to_' + self.to_now + '_sap_natur.XLS'))        # SAP auswertekategorien
        # Hiebssatzbilanz
        dat_hs = os.path.isfile(os.path.join(self.data_path, 'TO' + self.to_now, 'to_' + self.to_now + '_hs_bilanz.XLS'))           # HS-Bilanz neu
        dat_es = os.path.isfile(os.path.join(self.data_path, 'TO' + self.to_now, 'to_' + self.to_now + '_hs_bilanz_old.XLS'))       # HS-Bilanz alt
        # Berichtswesen
        dat_es_hs_1 = os.path.isfile(os.path.join(self.data_path, 'TO' + self.to_now, 'to_' + self.to_now + '_bw_alter.xlsx'))      # BW alter
        dat_es_hs_2 = os.path.isfile(os.path.join(self.data_path, 'TO' + self.to_now, 'to_' + self.to_now + '_bw_neig.xlsx'))       # BW neigung
        dat_es_hs_3 = os.path.isfile(os.path.join(self.data_path, 'TO' + self.to_now, 'to_' + self.to_now + '_bw_seeh.xlsx'))       # BW seehoehe
        dat_es_hs_4 = os.path.isfile(os.path.join(self.data_path, 'TO' + self.to_now, 'to_' + self.to_now + '_bw_uz.xlsx'))         # BW umtriebszeit
        dat_es_hs = dat_es_hs_1 & dat_es_hs_2 & dat_es_hs_3 & dat_es_hs_4
        dat_zv_ze = os.path.isfile(os.path.join(self.data_path, 'TO' + self.to_now, 'to_' + self.to_now + '_bw_zufaellige.xlsx'))   # BW zufällige
        dat_wpplan = os.path.isfile(os.path.join(self.data_path, 'TO' + self.to_now, 'to_' + self.to_now + '_bw_wp.xlsx'))          # BW waldpflegeplan
        # SPI
        dat_spi = os.path.isfile(os.path.join(self.data_path, 'stichprobe', 'SPI_' + self.fe_year + '.txt'))               # SPI
        # Klima
        dat_kima_1 = os.path.isfile(os.path.join(self.data_path, 'klima', '1_Lufttemperatur.txt'))                                  # Klima
        dat_kima_2 = os.path.isfile(os.path.join(self.data_path, 'klima', '2_Niederschlag.txt'))
        dat_kima_3 = os.path.isfile(os.path.join(self.data_path, 'klima', '4_Schnee.txt'))
        dat_kima = dat_kima_1 & dat_kima_2 & dat_kima_3

        #-----------------------------------------------------------------------------------------
        #--- get data
        #-----------------------------------------------------------------------------------------


        if dat_hs == True:
            path_hs = path_dir + '/to_' + self.to_now + '_hs_bilanz.XLS'
            obf_hs = OBFHiebssatz(pd.read_csv(path_hs, sep='\t', encoding = "ISO-8859-1", decimal=',', error_bad_lines=False))

        if dat_es == True:
            path_es_1 = path_dir + '/to_' + self.to_now + '_hs_bilanz_old.XLS'
            es_all = pd.read_csv(path_es_1, sep='\t', encoding = "ISO-8859-1", decimal=',', error_bad_lines=False)
            obf_es = OBFEinschlag(es_all)

        if dat_es_hs == True:
            path_alter = path_dir + '/to_' + self.to_now + '_bw_alter.xlsx'
            path_neig = path_dir + '/to_' + self.to_now + '_bw_neig.xlsx'
            path_seeh = path_dir + '/to_' + self.to_now + '_bw_seeh.xlsx'
            path_uz = path_dir + '/to_' + self.to_now + '_bw_uz.xlsx'
            obf_es_hs = OBFEinschlagHiebsatz(pd.read_excel(path_alter), pd.read_excel(path_neig), pd.read_excel(path_seeh), pd.read_excel(path_uz))

        if dat_zv_ze == True:
            path_zv_ze = path_dir + '/to_' + self.to_now + '_bw_zufaellige.xlsx'
            obf_zv_ze = OBFSchadholz(pd.read_excel(path_zv_ze))

        if dat_wpplan == True:
            path_wp = path_dir + '/to_' + self.to_now + '_bw_wp.xlsx'
            obf_wp = OBFWaldpflege(pd.read_excel(path_wp))

        if dat_natur == True:
            path_natur = path_dir + '/to_' + self.to_now + '_sap_natur.XLS'
            obf_natur = OBFNatur(pd.read_csv(path_natur, sep='\t', encoding = "ISO-8859-1", decimal=',', error_bad_lines=False))
            # set dictionary for natur
            obf_natur.set_dic(path_dict + '/auswertekat_index.xlsx')

        if dat_spi == True:
            path_spi = os.path.join(self.data_path, 'stichprobe', 'SPI_' + self.fe_year + '.txt')
            obf_spi = OBFSpi(pd.read_csv(path_spi, sep='\t', encoding = "ISO-8859-1", decimal=',', thousands='.', skipinitialspace=True, skiprows=3))

        if dat_kima == True:
            path_klima = os.path.join(self.data_path, 'klima')
            obf_klima = OBFKlima(pd.read_csv(path_klima + '/1_Lufttemperatur.txt', sep='\t', header=0), pd.read_csv(path_klima + '/2_Niederschlag.txt', sep='\t', header=0), pd.read_csv(path_klima + '/4_Schnee.txt', sep='\t', header=0))

        obf_text = OBFText()
        obf_dic = OBFDictionary()

        obf_et = OBFEt(pd.read_excel(path_dict + '/Ertragstafelsets.xlsx', 'Baumarten'), obf_fuc.data[['Betriebsklasse', 'Ertragstafelnummer', 'Baumart_group']])

        # set legendes
        obf_text.set_legend(pd.read_excel(path_dict + '/key.xlsx', 'Dringlichkeit'), pd.read_excel(path_dict + '/key.xlsx', 'Bewilligung'), pd.read_excel(path_dict + '/key.xlsx', 'Nutzzeitpunkt'), pd.read_excel(path_dict + '/key.xlsx', 'Schlägerungsart'), pd.read_excel(path_dict + '/key.xlsx', 'Rückungsart'), pd.read_excel(path_dict + '/key.xlsx', 'Massnahmenart'), pd.read_excel(path_dict + '/key.xlsx', 'Exposition'), pd.read_excel(path_dict + '/key.xlsx', 'Standort'), pd.read_excel(path_dict + '/key.xlsx', 'ZE'), pd.read_excel(path_dict + '/key.xlsx', 'ZV'))

        log = "0.1 load data - successful\n"
        self.save_log(log)

        # create old to text

        self.to_old_text =''

        for i in self.to_old:
            self.to_old_text = self.to_old_text + str(i) + ', '

        self.to_old_text = self.to_old_text[:-2]

        obf_doc = OBFDocX(Document(path_dict + '/templet_xx.docx'))

        #######################################################################################################################
        ###   0 Deckblatt
        #######################################################################################################################

        print('***   0 Deckblatt   ***')

        obf_doc.doc.add_picture(path_dict + '/oebf.png')

        for i in range(4):
            obf_doc.doc.add_paragraph('')

        p = obf_doc.doc.add_paragraph().add_run('OPERAT')
        p.bold = True
        p.font.name = 'Verdana'
        p.font.size = Pt(34)

        p = obf_doc.doc.add_paragraph().add_run('FB ' + str(obf_fuc.dic.fb) + ' ' + obf_fuc.dic.dic_num_fb[obf_fuc.dic.fb])
        p.bold = True
        p.font.name = 'Verdana'
        p.font.size = Pt(35)

        obf_doc.doc.add_paragraph('')

        p = obf_doc.doc.add_paragraph().add_run('Teiloperat ' + str(obf_fuc.dic.to))
        p.font.name = 'Verdana'
        p.font.size = Pt(25)

        for i in range(3):
            obf_doc.doc.add_paragraph('')

        for i in obf_fuc.loop_fr()[1:]:

            p = obf_doc.doc.add_paragraph().add_run(obf_fuc.name_code(i)[1])
            p.font.name = 'Verdana'
            p.font.size = Pt(20)

        for i in range(12-(obf_fuc.loop_fr()[1:,0].shape[0]*2)):
            p = obf_doc.doc.add_paragraph().add_run('')
            p.font.size = Pt(20)

        p = obf_doc.doc.add_paragraph().add_run('Organisationsstand 01.01.' + obf_fuc.dic.laufzeit_begin)
        p.font.name = 'Verdana'
        p.font.size = Pt(18)

        p = obf_doc.doc.add_paragraph().add_run('Laufzeit 01.01.' + obf_fuc.dic.laufzeit_begin + ' - 31.12.' + obf_fuc.dic.laufzeit_end)
        p.font.name = 'Verdana'
        p.font.size = Pt(18)

        for i in range(10-(obf_fuc.loop_fr()[1:,0].shape[0]*2)):
            obf_doc.doc.add_paragraph('')

        p = obf_doc.doc.add_paragraph().add_run('Für die Unternehmensleitung                       Für den Forstbetrieb')
        p.font.name = 'Verdana'
        p.font.size = Pt(12)

        obf_doc.doc.add_page_break()

        for i in range(3):
            obf_doc.doc.add_paragraph('')

        run = obf_doc.doc.add_paragraph('Erstellung vom Teiloperat mit AutOP ').add_run(' Version ' + version)
        font = run.font
        font.name = 'Verdana'
        font.size = Pt(8)

        obf_doc.doc.add_paragraph('')
        obf_doc.doc.add_paragraph('Anleitung zum Fertigstellen des Operates')
        obf_doc.doc.add_paragraph('')
        obf_doc.doc.add_paragraph('Die Datenaufarbeitung für das Operat erfolgt mit Hilfe des Programms AutOP. Dieses Programm verarbeitet die SAP-Daten und erstellt ein größtenteils fertiges Operat im Word-Format. Einzelne individuelle Textblöcke, komplexere Tabellen und Formatierungen müssen dennoch händisch ergänzt werden. Folgende Liste soll einen Überblick geben, welche Teile noch ausgearbeitet gehören. Zusätzlich sind alle zu bearbeitenden Textstellen, Tabellen und Abbildungen mit')
        obf_doc.doc.add_paragraph('***   feier Text   ***')
        obf_doc.doc.add_paragraph('***   xTabelle   ***')
        obf_doc.doc.add_paragraph('***   xAbbildung   ***')
        obf_doc.doc.add_paragraph('hervorgehoben.')

        obf_doc.doc.add_paragraph('')
        obf_doc.doc.add_paragraph('In folgenden Kapiteln gehören die fehlenden Informationen/Tabellen ergänzt:')
        obf_doc.doc.add_paragraph('1.1 Einteilung der Taxation (Tabelle)')
        obf_doc.doc.add_paragraph('2.3 Einforstung (Tabellen)')
        obf_doc.doc.add_paragraph('3.7 Herkunftsgebiete und Höhenstufen (Tabelle)')
        obf_doc.doc.add_paragraph('4.x Besitzstand (ganzes Kapitel)')
        obf_doc.doc.add_paragraph('6.1.4 Waldbauliche Beurteilung - Nutzungen (Text)')
        obf_doc.doc.add_paragraph('6.2.2 Berechneter Hiebssatz (Tabellen / Grafiken aus Hiebssatzbegründung)')
        obf_doc.doc.add_paragraph('6.2.4 Hiebssatz Begründung (Text aus Hiebssatzbegründung)')
        obf_doc.doc.add_paragraph('6.2.5 Vergleich Hiebssatz vergangene und aktuelle Periode (Tabelle aus Hiebssatzbegründung)')
        obf_doc.doc.add_paragraph('7.1.1 Waldbauliche Beurteilung - Waldpflege (Text)')
        obf_doc.doc.add_paragraph('7.2 Baumartenwahl (Tabelle)')
        obf_doc.doc.add_paragraph('8.x Einforstung (ganzes Kapitel)')
        obf_doc.doc.add_paragraph('11.x Vormerkungen für die nächste Forsteinrichtung (Text)')

        obf_doc.doc.add_paragraph('')
        obf_doc.doc.add_paragraph('Die Formatierung in folgenden Punkten muss händisch überarbeitet werden:')
        obf_doc.doc.add_paragraph('Versetzte Tabellen der allgemeinen Formatierung entsprechend verschieben (insbesondere im Kapitel 3. Natürliche Grundlagen)', style='List Bullet')
        obf_doc.doc.add_paragraph('Bei Tabellen, die nicht auf einer Seite Platz haben, die Titelzeile der Tabelle duplizieren', style='List Bullet')
        obf_doc.doc.add_paragraph('Ausrichtung einzelner Seiten ins Querformat ändern (insbesondre in den Kapiteln 2.1 Flächenübersicht, 3.5 Standorteinheiten, 7.1 Waldpflege)', style='List Bullet')
        obf_doc.doc.add_paragraph('Seitenumbrüche anlegen, sodass', style='List Bullet')
        obf_doc.doc.add_paragraph('   Tabellen, welche über 2 Seiten gehen, auf einer angezeigt werden', style='List Bullet')
        obf_doc.doc.add_paragraph('   Kapitelüberschriften am Seitenanfang stehen', style='List Bullet')

        obf_doc.doc.add_paragraph('')
        obf_doc.doc.add_paragraph('Abschließend müssen die Verzeichnisse aktualisiert werden:')
        obf_doc.doc.add_paragraph('Erstellen vom Inhaltsverzeichnis')
        obf_doc.doc.add_paragraph('Erstellen vom Tabellenverzeichnis')
        obf_doc.doc.add_paragraph('Erstellen vom Abbildungsverzeichnis')

        obf_doc.doc.add_page_break()

        log = log + "0.2 Deckblatt - successful\n"
        self.save_log(log)


        #######################################################################################################################
        ###   1 Allgemeines
        #######################################################################################################################

        if self.do_sections['1_Allgemeines'] == 1:
            try:

                print('***   1 Allgemeines   ***')
                obf_doc.doc.add_heading('Allgemeines', 1)

                obf_doc.doc.add_heading('Erstellung und Erfassung', 2)

                # add paragraph
                obf_doc.docx_paragraph_table('Einteilung der Taxation')
                obf_doc.doc.add_paragraph('*** xTabelle ***')

                obf_doc.doc.add_heading('Verwendete Ertragstafeln', 2)

                et_sets = obf_et.fuc_tbl_et_set()

                # add paragraph
                obf_doc.docx_paragraph_table('Zuordnung Ertragstafelset zu Betriebsklassen')
                # add table
                obf_doc.docx_table_text(et_sets, width=[3,3,10], header_rep = True, header = 'Ertragstafelset in den Betriebsklassen', font_size = 7, autofit = True)
                obf_doc.doc.add_paragraph('')

                for sets in et_sets['Ertragstafelset'].unique():

                    table = obf_et.fuc_tbl_et_ba(sets, allba=False)

                    # add paragraph
                    obf_doc.docx_paragraph_table('Zuordnung der Ertragstafeln zu den Baumarten im Ertragstafelset ' + str(sets) + '' + '')

                    # add table
                    obf_doc.docx_table_text(table, width=[2,2,5,5,2], header_rep = True, header = 'Ertragstafeln', font_size = 7, autofit = True)
                    obf_doc.doc.add_paragraph('')

                obf_doc.doc.add_page_break()
            except:
                log = log + "1.0 Allgemeines - error occured\n"
                self.save_log(log)
            else:
                log = log + "1.0 Allgemeines - successful\n"
                self.save_log(log)


        #######################################################################################################################
        ###   2 Hauptergebnisse
        #######################################################################################################################

        if self.do_sections['2_Hauptergebnisse'] == 1:

            # add section break
            new_section = obf_doc.doc.add_section(WD_SECTION.NEW_PAGE)
            new_section.start_type

            print('***   2 Hauptergebnisse   ***')
            obf_doc.doc.add_heading('Hauptergebnisse', 1)

            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   2.1 Flächenübersicht
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            try:
                print('Flächenübersicht')

                obf_doc.doc.add_heading('Flächenübersicht', 2)

                obf_doc.doc.add_heading('Allgemein', 3)

                table = obf_fuc.fuc_tbl_flaechen()

                # add paragraph
                obf_doc.docx_paragraph_table('Flächenübersicht FB ' + obf_fuc.dic.dic_num_fb[obf_fuc.dic.fb] + ', Teiloperat ' + str(obf_fuc.dic.to))

                # add table
                x = obf_doc.get_x('flaechenuebersicht')
                obf_doc.docx_table_3x(table, x, header_rep = True, header = 'Flächenübersicht [ha]', font_size = 7, autofit = True)
                obf_doc.doc.add_paragraph('')

                # if schutzwald existing, otherwise no print
                if obf_fuc.sw:

                    obf_doc.doc.add_heading('Schutzwald', 3)
                    table = obf_fuc.fuc_tbl_flaechen_sw()

                    # add paragraph
                    obf_doc.docx_paragraph_table('Flächenverteilung der Schutzwaldkategorien im FB ' + obf_fuc.dic.dic_num_fb[obf_fuc.dic.fb] + ', Teiloperat ' + str(obf_fuc.dic.to))

                    # add table
                    x = obf_doc.get_x('sw_flaechenuebersicht')
                    obf_doc.docx_table_3x(table, x, header_rep = True, header = 'Schutzwald Flächenübersicht [ha]', font_size = 7, autofit = True)
                    obf_doc.doc.add_paragraph('')
                    #obf_doc.docx_table(table, Cm(2.2), Cm(2.2))

                # add section break
                new_section = obf_doc.doc.add_section(WD_SECTION.NEW_PAGE)
                new_section.start_type
            except:
                log = log + "2.1 Flächenübersicht - error occured\n"
                self.save_log(log)
            else:
                log = log + "2.1 Flächenübersicht - successful\n"
                self.save_log(log)

            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   2.2 Hiebssätze im Jahrzehnt
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###

            if dat_hs == True:
                try:

                    print('Hiebssätze im Jahrzehnt ' + obf_fuc.dic.laufzeit)

                    obf_doc.doc.add_heading('Hiebssätze im Jahrzehnt ' + obf_fuc.dic.laufzeit, 2)

                    print('Hiebssätze Forstrevier')
                    table = obf_hs.fuc_tbl_hs_fest('Forstrevier')

                    # add paragraph
                    obf_doc.docx_paragraph_table('Festgesetzter dezennaler Hiebssatz im FB ' + obf_fuc.dic.dic_num_fb[obf_fuc.dic.fb] + ', Teiloperat ' + str(obf_fuc.dic.to) + ' ' + obf_fuc.dic.laufzeit)

                    # add table
                    x = obf_doc.get_x('festgesetzterHS')
                    obf_doc.docx_table_3x(table, x, header_rep = True, header = 'Festgesetzte Hiebssätze [Efm]', font_size = 7, autofit = True)
                    obf_doc.doc.add_paragraph('')
                except:
                    log = log + "2.2 Hiebssätze im Jahrzehnt - error occured\n"
                    self.save_log(log)
                else:
                    log = log + "2.2 Hiebssätze im Jahrzehnt - successful\n"
                    self.save_log(log)


            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   2.3 Einforstung
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            try:
                # add section break
                new_section = obf_doc.doc.add_section(WD_SECTION.NEW_PAGE)
                new_section.start_type

                print('Einforstung')

                obf_doc.doc.add_heading('Einforstung', 2)

                obf_doc.doc.add_heading('Holz- und Streubezugsrechte', 3)
                # add paragraph
                #obf_doc.docx_paragraph_table('Holz- und Streubezugsrechte im FB ' + obf_fuc.dic.dic_num_fb[obf_fuc.dic.fb] + ', Teiloperat ' + str(obf_fuc.dic.to))
                obf_doc.doc.add_paragraph('*** Tabelle falls Flächen einforstet sind ***')

                obf_doc.doc.add_heading('Weiderechte', 3)
                # add paragraph
                #obf_doc.docx_paragraph_table('Weiderechte im FB ' + obf_fuc.dic.dic_num_fb[obf_fuc.dic.fb] + ', Teiloperat ' + str(obf_fuc.dic.to))
                obf_doc.doc.add_paragraph('*** Tabelle falls Flächen mit Weiderechten vohanden sind ***')

                # add section break
                new_section = obf_doc.doc.add_section(WD_SECTION.NEW_PAGE)
                new_section.start_type
                #obf_doc.doc.add_page_break()
            except:
                log = log + "2.3 Einforstung - error occured\n"
                self.save_log(log)
            else:
                log = log + "2.3 Einforstung - successful\n"
                self.save_log(log)

            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   2.4 Fläche und Vorrat nach Altersklassen
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            try:
                print('Fläche und Vorrat nach Altersklassen')

                obf_doc.doc.add_heading('Fläche und Vorrat nach Altersklassen', 2)

                for i in obf_fuc.loop_be():

                    print(i)
                    table = obf_fuc.fuc_tbl_plt_flaeche_vorrat(i)

                    # add figure
                    obf_doc.doc.add_picture('tempx.png')

                    # add Caption to figure
                    obf_doc.docx_paragraph_figure('Fläche und Vorrat nach Altersklassen im FB ' + obf_fuc.dic.dic_num_fb[obf_fuc.dic.fb] + ', Teiloperat ' + str(obf_fuc.dic.to) + ', ' + obf_fuc.name_code(i)[3] + ' (AKL VIII+ enthält die 8. Altersklasse und älter sowie die Überhälter)')

                    # add Caption to table
                    obf_doc.docx_paragraph_table('Fläche und Vorrat nach Altersklassen im FB ' + obf_fuc.dic.dic_num_fb[obf_fuc.dic.fb] + ', Teiloperat ' + str(obf_fuc.dic.to) + ', ' + obf_fuc.name_code(i)[3] + ' (AKL VIII+ enthält die 8. Altersklasse und älter sowie die Überhälter)')

                    # add table
                    obf_doc.docx_table_x(table, Cm(2.2), Cm(2.2), False, '', font_size = 7, autofit = True)
                    obf_doc.doc.add_paragraph('')

                    obf_doc.doc.add_page_break()
            except:
                log = log + "2.4 Fläche und Vorrat nach Altersklassen - error occured\n"
                self.save_log(log)
            else:
                log = log + "2.4 Fläche und Vorrat nach Altersklassen - successful\n"
                self.save_log(log)

            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   2.5 Baumarten
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            try:
                print('Baumarten')

                obf_doc.doc.add_heading('Baumarten', 2)

                for i in obf_fuc.loop_be():

                    print(i)
                    table = obf_fuc.fuc_tbl_plt_ba_anteile(i)

                    # add figure
                    obf_doc.doc.add_picture('tempx.png')
                    # add Caption to figure
                    obf_doc.docx_paragraph_figure('Baumartenanteile in den Altersklassen nach Flächenanteilen im FB ' + obf_fuc.dic.dic_num_fb[obf_fuc.dic.fb] + ', Teiloperat ' + str(obf_fuc.dic.to) + ', ' + obf_fuc.name_code(i)[3] + ' (' + obf_fuc.dic.sl + obf_fuc.dic.sn + 'AKL VIII enthält die 8. Altersklasse und älter sowie die Überhälter)')

                    # add Caption to table
                    obf_doc.docx_paragraph_table('Baumartenanteile in den Altersklassen nach Flächenanteilen im FB ' + obf_fuc.dic.dic_num_fb[obf_fuc.dic.fb] + ', Teiloperat ' + str(obf_fuc.dic.to) + ', ' + obf_fuc.name_code(i)[3] + ' (' + obf_fuc.dic.sl + obf_fuc.dic.sn + 'AKL VIII enthält die 8. Altersklasse und älter sowie die Überhälter)')

                    # add table
                    obf_doc.docx_table_x(table, Cm(2.2), Cm(2.2), True, 'Baumartenanteile ' + obf_fuc.name_code(i)[0] + ' ' + obf_fuc.name_code(i)[2] + ' [%]', font_size = 7, autofit = True)
                    obf_doc.doc.add_paragraph('')

                    obf_doc.doc.add_page_break()
            except:
                log = log + "2.5 Baumarten - error occured\n"
                self.save_log(log)
            else:
                log = log + "2.5 Baumarten - successful\n"
                self.save_log(log)

            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   2.6 Mittlere Ertragsklasse der Hauptbaumarten
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            try:
                print('Mittlere Ertragsklassen der Hauptbaumarten')

                obf_doc.doc.add_heading('Mittlere Ertragsklassen der Hauptbaumarten', 2)

                table = obf_fuc.fuc_tbl_mittel_('Ertragsklasse')

                # add paragraph
                obf_doc.docx_paragraph_table('Mittlere Ertragsklassen der Baumarten im FB ' + obf_fuc.dic.dic_num_fb[obf_fuc.dic.fb] + ', Teiloperat ' + str(obf_fuc.dic.to))

                # add table
                obf_doc.docx_table_x(table, Cm(2.2), Cm(2.2), header_rep = True, header = 'Mittlere Ertragsklassen im TO ' + str(obf_fuc.dic.to) + ' [Vfm/ha/a]', font_size = 7, autofit = True)
                obf_doc.doc.add_paragraph('')

                obf_doc.doc.add_page_break()
            except:
                log = log + "2.6 Mittlere Ertragsklasse der Hauptbaumarten - error occured\n"
                self.save_log(log)
            else:
                log = log + "2.6 Mittlere Ertragsklasse der Hauptbaumarten - successful\n"
                self.save_log(log)


        #######################################################################################################################
        ###   3 Natürliche Grundlagen
        #######################################################################################################################

        if self.do_sections['3_Natuerliche_Grundlagen'] == 1:

            print('***   3 Natürliche Grundlagen   ***')
            obf_doc.doc.add_heading('Natürliche Grundlagen', 1)

            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   3.1 Seehöhe | 3.2 Exposition | 3.3 Neigung
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            try:
                for j in ['Seehöhe', 'Exposition', 'NeigGruppe']:

                    print(j)

                    obf_doc.doc.add_heading(obf_fuc.name_ng(j)[1], 2)

                    table_ng = obf_fuc.fuc_tbl_natuerliche_grundlagen(j)

                    for i in table_ng.columns:
                        print(i)

                        table = obf_fuc.fuc_plt_natuerliche_grundlagen(table_ng,j,i,obf_fuc.name_ng(j,i))

                        # add figure
                        obf_doc.doc.add_picture('tempx.png')

                        # add table
                        obf_doc.docx_table_x(table, Cm(1.2), Cm(1.2), header_rep = True, header = obf_fuc.name_ng(j,i)[3], font_size = 7, autofit = True)
                        obf_doc.doc.add_paragraph('')

                    # add Caption to figure
                    obf_doc.docx_paragraph_figure(obf_fuc.name_ng(j)[0] + ' der Waldorte in den Forstrevieren und im Teiloperat')

                    obf_doc.doc.add_page_break()
            except:
                log = log + "3.1 Seehöhe | 3.2 Exposition | 3.3 Neigung - error occured\n"
                self.save_log(log)
            else:
                log = log + "3.1 Seehöhe | 3.2 Exposition | 3.3 Neigung - successful\n"
                self.save_log(log)

            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   3.4 Klima
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###

            print('Klima')

            # add heading
            obf_doc.doc.add_heading('Klima', 2)
            obf_doc.doc.add_paragraph('*** Text, Tabellen, Grafiken zum Klima ***')
            obf_doc.doc.add_paragraph('')

            # # # check
            if dat_kima == True:
                try:

                    stations = self.klima_stationen

                    # add heading
                    obf_doc.doc.add_heading('Klimadaten', 3)
                    print('Klimadiagramme')

                    obf_doc.docx_paragraph_table('Klimanormalwerte (Periode 1981-2010) für relevante Stationen im Bereich des TO, Tage mit min. 1 cm Schnee = Mittlere Anzahl der Tage mit mindestens 1 cm Schneedeckenhöhe, Tage mit min. 20 cm Schnee = Mittlere Anzahl der Tage mit mindestens 20 cm Schneedeckenhöhe (Datenquelle: ZAMG, 2015)')
                    table = obf_klima.fuc_tbl_climate(stations)
                    obf_doc.docx_table_x(table, Cm(1.2), Cm(1.2), header_rep = True, header = 'Klimadaten', font_size = 7, autofit = True)
                    obf_doc.doc.add_paragraph('')

                    # add heading
                    obf_doc.doc.add_heading('Klimadiagramme', 3)
                    print('Klimadiagramme')

                    obf_klima.fuc_plt_climate(stations, 'temp')
                    obf_doc.doc.add_picture('tempx.png')
                    # add Caption to figure
                    obf_doc.docx_paragraph_figure('Monatsmitteltemperaturen (Periode 1981-2010) für relevante Stationen im Bereich des TO ' + str(obf_fuc.dic.to) + ', (Datenquelle: ZAMG 2015)')

                    obf_klima.fuc_plt_climate(stations, 'prec_month')
                    obf_doc.doc.add_picture('tempx.png')
                    # add Caption to figure
                    obf_doc.docx_paragraph_figure('Mittlerer Monatsniederschlag (Periode 1981-2010) für relevante Stationen im Bereich des TO ' + str(obf_fuc.dic.to) + ', (Datenquelle: ZAMG 2015)')

                    obf_klima.fuc_plt_climate(stations, 'prec_day')
                    obf_doc.doc.add_picture('tempx.png')
                    # add Caption to figure
                    obf_doc.docx_paragraph_figure('Maximaler monatlicher Tagesniederschlag (Periode 1981-2010) für relevante Stationen im Bereich des TO ' + str(obf_fuc.dic.to) + ', (Datenquelle: ZAMG 2015)')

                    obf_klima.fuc_plt_climate_snow(stations)
                    obf_doc.doc.add_picture('tempx.png')
                    # add Caption to figure
                    obf_doc.docx_paragraph_figure('Anzahl der Tage mit geschlossener Schneedecke (Periode 1981-2010) für relevante Stationen im Bereich des TO ' + str(obf_fuc.dic.to) + '; min 1cm = Mittlere Anzahl der Tage mit mindestens 1 cm Schneedeckenhöhe; min 20cm = Mittlere Anzahl der Tage mit mindestens 20 cm Schneedeckenhöhe; (Datenquelle: ZAMG 2015)')

                    obf_doc.doc.add_page_break()
                except:
                    log = log + "3.4 Klima - eerror occured\n"
                    self.save_log(log)
                else:
                    log = log + "3.4 Klima - successful\n"
                    self.save_log(log)

            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   3.5 Standort
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            try:
                print('Standort')

                # add section break
                new_section = obf_doc.doc.add_section(WD_SECTION.NEW_PAGE)
                new_section.start_type

                # add heading
                obf_doc.doc.add_heading('Standort', 2)

                # add heading
                obf_doc.doc.add_heading('Standortseinheiten im Teiloperat', 3)

                #obf_doc.docx_table_leg(leg_stoe, Cm(1), Cm(5))

                # get unique Standorteinheiten
                stoe_unq = obf_fuc.data_stoe['Standorteinheit'].unique()
                # load Standorteinheiten-Info
                stoe_info = pd.read_excel(path_dict + '/3_data_Standortseinheiten.xlsx')
                # filter just the unique Standorteinheiten
                stoe_info = stoe_info[stoe_info['STO-NR'].isin(stoe_unq)]
                # fill nan with -
                stoe_info = stoe_info.fillna('-')

                # add Caption to table
                obf_doc.docx_paragraph_table('Übersicht der Standortseinheiten im FB ' + obf_fuc.dic.dic_num_fb[obf_fuc.dic.fb] + ', Teiloperat ' + str(obf_fuc.dic.to) + '; sortiert nach Nummer der Standortseinheit. Für weitere Informationen siehe ÖBf Waldbauhandbuch.')

                # add table
                obf_doc.docx_table_stoe(stoe_info, Cm(1), Cm(2.2))
                obf_doc.doc.add_paragraph('')

                obf_doc.doc.add_page_break()


                ###   Anteile der Standortseinheiten_hektar   ###
                print('   Anteile der Standortseinheiten_hektar')

                table = obf_fuc.fuc_tbl_plt_standortseinheiten()

                # add figure
                obf_doc.doc.add_picture('tempx.png')
                # add Caption to figure
                obf_doc.docx_paragraph_figure('Anteile der Standortseinheiten in den Forstrevieren und im Teiloperat ' + str(obf_fuc.dic.to) + ' in Prozent, die Bedeutung der Standortseinheiten ist im Kapiltel 3.5.1 Standortseinheiten im Teiloperat ersichtilich')

                # add Caption to table
                obf_doc.docx_paragraph_table('Flaäche der Standortseinheiten in den Forstrevieren und im Teiloperat ' + str(obf_fuc.dic.to) + ' in Hektar, die Bedeutung der Standortseinheiten ist im Kapiltel 3.5.1 Standortseinheiten im Teiloperat ersichtilich')

                # add table
                obf_doc.docx_table_x(table, Cm(2), Cm(1), header_rep = False, header = '', font_size = 7, autofit = True)
                obf_doc.doc.add_paragraph('')

                # add section break
                new_section = obf_doc.doc.add_section(WD_SECTION.NEW_PAGE)
                new_section.start_type


                ###   Substrat und Wüchsigkeit im Teiloperat   ###
                print('   Substrat und Wüchsigkeit im Teiloperat')

                # add heading
                obf_doc.doc.add_heading('Substrat und Wüchsigkeit im Teiloperat', 3)

                for i in obf_fuc.loop_fr():

                    table = obf_fuc.fuc_tbl_plt_geologie_wuechsigkeit(i)

                    # add figure
                    obf_doc.doc.add_picture('tempx.png')
                    # add Caption to figure
                    obf_doc.docx_paragraph_figure('Übersicht über Substrat und Wüchsigkeit im ' + obf_fuc.name_code(i)[1])
                    # add Caption to table
                    obf_doc.docx_paragraph_table('Flächenverteilung der Wüchsigkeitsgruppen nach dem Substrat im ' + obf_fuc.name_code(i)[1])
                    # add table
                    obf_doc.docx_table_x(table, Cm(2.2), Cm(2.2), header_rep = True, header = 'Wüchsigkeitsgruppen im ' + obf_fuc.name_code(i)[1] + ' [ha]', font_size = 7, autofit = True)
                    obf_doc.doc.add_paragraph('')

                obf_doc.doc.add_page_break()
            except:
                log = log + "3.5 Standort - error occured\n"
                self.save_log(log)
            else:
                log = log + "3.5 Standort - successful\n"
                self.save_log(log)

            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   3.6 Wuchsgebiete
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            try:
                print('Wuchsgebiete')

                obf_doc.doc.add_heading('Wuchsgebiete', 2)

                obf_doc.doc.add_picture(path_dict + '/3_wuchsgebiete.png')

                # add Caption to figure
                obf_doc.docx_paragraph_figure('Forstliche Wuchsgebiete Österreichs (Kilian et al., 1994)')

                # add Caption to table
                obf_doc.docx_paragraph_table('Gliederung der forstlichen Wuchsgebiete Österreichs. Die Gliederung umfasst 22 Wuchsgebiete, die in 9 Hauptwuchsgebiete zusammengefasst sind (Kilian et al., 1994).')

                wg_info = pd.read_excel(path_dict + '/3_data_Wuchsgebiete.xlsx')

                obf_doc.docx_table_text(wg_info, width=[7,0.5,8], header_rep = True, header = 'Forstliche Wuchsgebiete Österreichs', font_size = 7, autofit = True)
                obf_doc.doc.add_paragraph('')

                # add text for the used Wuchsgebiete (ex. 4.2) from Wuchsgebiete Österreichs Kilian et. al.
                for i in obf_fuc.dic.wg:
                    # insert text
                    obf_text.fuc_txt_wuchsgebiet(obf_doc.doc, i)

                #obf_doc.doc.add_page_break()
            except:
                log = log + "3.6 Wuchsgebiete - error occured\n"
                self.save_log(log)
            else:
                log = log + "3.6 Wuchsgebiete - successful\n"
                self.save_log(log)

            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   3.7 Herkunftsgebiete und Höhenstufen
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            try:
                print('Herkunftsgebiete und Höhenstufen')

                # add section break
                new_section = obf_doc.doc.add_section(WD_SECTION.NEW_PAGE)
                new_section.start_type

                obf_doc.doc.add_heading('Herkunftsgebiete und Höhenstufen', 2)

                obf_doc.docx_paragraph_table('Herkunftsgebiete und Höhenstufen nach Kilian et al. 1994. Für das Teiloperat relevante Herkunftsgebiete wurden orange hervorgehoben.')
                obf_doc.doc.add_paragraph('*** xTabelle ***')

                # add section break
                new_section = obf_doc.doc.add_section(WD_SECTION.NEW_PAGE)
                new_section.start_type
            except:
                log = log + "3.7 Herkunftsgebiete und Höhenstufen - error occured\n"
                self.save_log(log)
            else:
                log = log + "3.7 Herkunftsgebiete und Höhenstufen - successful\n"
                self.save_log(log)


        #######################################################################################################################
        ###   4 Besitzstand
        #######################################################################################################################

        if self.do_sections['4_Besitzstand'] == 1:
            try:

                print('***   4 Besitzstand   ***')
                obf_doc.doc.add_heading('Besitzstand', 1)

                obf_doc.doc.add_heading('Flächenübersicht', 2)
                obf_doc.docx_paragraph_table('Flächenveränderungen im Teiloperat von der vergangenen zur aktuellen Forsteirichtungsperiode')
                obf_doc.doc.add_paragraph('*** xTabelle ***')
                obf_doc.doc.add_paragraph('')
                obf_doc.doc.add_paragraph('*** freier Text ***')
                obf_doc.doc.add_paragraph('')

                obf_doc.doc.add_heading('Grundverkehr', 2)
                obf_doc.doc.add_paragraph('*** freier Text ***')
                obf_doc.doc.add_paragraph('')

                obf_doc.doc.add_heading('Organisationsänderungen', 2)
                obf_doc.doc.add_paragraph('*** freier Text ***')
                obf_doc.doc.add_paragraph('')

                obf_doc.doc.add_heading('Änderungen bei den Teiloperaten', 3)
                obf_doc.doc.add_paragraph('*** freier Text ***')
                obf_doc.doc.add_paragraph('')

                obf_doc.doc.add_heading('Betriebsklassen – Umschlüsselung', 2)
                obf_doc.doc.add_paragraph('*** freier Text ***')
                obf_doc.doc.add_paragraph('')

                obf_doc.doc.add_heading('Änderungen der räumlichen Einteilung (RET)', 2)
                obf_doc.doc.add_paragraph('Grundsätzlich wurden die Einteilungslinien im ursprünglichen Zustand belassen. Gelegentlich wurden die Einteilungslinien nach Möglichkeit auf Straßen, Rücken, Bäche, oder andere in der Natur leicht auffindbare Linien verlegt.')
                obf_doc.doc.add_paragraph('In der Natur nicht erkennbare Bestandesgrenzen wurden aufgelöst. Kleinstbestände, die keine eigenständige Bewirtschaftungseinheit darstellen, wurde zu größeren Einheiten zusammengefasst. Deren Unterschiede wurden im Text der Bestandesbeschreibungen festgehalten.')
                obf_doc.doc.add_paragraph('')

                obf_doc.doc.add_heading('Zustand der Versteinung', 2)
                obf_doc.doc.add_paragraph('*** freier Text ***')
                obf_doc.doc.add_paragraph('')

                obf_doc.doc.add_heading('Flächentabelle', 2)
                obf_doc.doc.add_paragraph('Die Flächentabelle ist in einem eigenem Buch dem Operat beigelegt.')
                obf_doc.doc.add_paragraph('')

                obf_doc.doc.add_heading('Erschließung', 2)
                obf_doc.docx_paragraph_table('Stand der Erschließung in den Forstrevieren und im Teiloperat.')
                obf_doc.doc.add_paragraph('*** xTabelle ***')
                obf_doc.doc.add_paragraph('')
                obf_doc.doc.add_paragraph('*** freier Text ***')

                obf_doc.doc.add_page_break()
            except:
                log = log + "4.0 Besitzstand - error occured\n"
                self.save_log(log)
            else:
                log = log + "4.0 Besitzstand - successful\n"
                self.save_log(log)


        #######################################################################################################################
        ###   5 Wald
        #######################################################################################################################

        if self.do_sections['5_Wald'] == 1:

            print('***   5 Wald   ***')
            obf_doc.doc.add_heading('Wald', 1)

            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   5.1 Fläche und Vorrat nach Altersklassen
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            try:
                print('Fläche und Vorrat nach Altersklassen')
                obf_doc.doc.add_heading('Fläche und Vorrat nach Altersklassen', 2)

                code = obf_fuc.loop_fr_uz()

                for i in code:

                    if np.array_equal(np.array(i), code[0]):
                        obf_doc.doc.add_heading('Revierweise', 3)
                    elif (i[0] == '0') & (int(i[1]) == obf_fuc.get_min_uz(code)):
                        obf_doc.doc.add_heading('Umtriebszeit', 3)

                    print(i)
                    table = obf_fuc.fuc_tbl_plt_flaeche_vorrat(i)

                    # add figure
                    obf_doc.doc.add_picture('tempx.png')

                    # add Caption to figure
                    obf_doc.docx_paragraph_figure('Fläche und Vorrat nach Altersklassen ' + obf_fuc.name_code(i)[1] + ', '  + obf_fuc.name_code(i)[3] + ' (AKL VIII+ enthält die 8. Altersklasse und älter sowie die Überhälter)')

                    # add Caption to table
                    obf_doc.docx_paragraph_table('Fläche und Vorrat nach Altersklassen ' + obf_fuc.name_code(i)[1] + ', '  + obf_fuc.name_code(i)[3] + ' (AKL VIII+ enthält die 8. Altersklasse und älter sowie die Überhälter)')

                    # add table
                    obf_doc.docx_table_x(table, Cm(2.2), Cm(2.2), False, '', 7, True)
                    obf_doc.doc.add_paragraph('')

                obf_doc.doc.add_page_break()
            except:
                log = log + "5.1 Fläche und Vorrat nach Altersklassen - error occured\n"
                self.save_log(log)
            else:
                log = log + "5.1 Fläche und Vorrat nach Altersklassen - successful\n"
                self.save_log(log)

            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   5.2 Vergleich der Vorräte im Wirtschaftswald, Taxation und Stichprobeninventur
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###

            if dat_spi == True:
                try:

                    print('Stichprobeninventur')
                    obf_doc.doc.add_heading('Vergleich der Vorräte im Wirtschaftswald, Taxation und Stichprobeninventur', 2)

                    ## Allgemeines
                    obf_doc.doc.add_heading('Stichprobeninventur',3)

                    obf_doc.doc.add_paragraph('Die Stichprobeninventur (SPI) erfolgte zeitgleich mit der Erhebung der Taxationsdaten. Die SPI wurde jedoch entkoppelt von der Taxation von Studenten durchgeführt und von DI Robert Zeiner ausgewertet.')
                    obf_doc.doc.add_paragraph('')
                    obf_doc.doc.add_paragraph('Die SPI ist als permanente Stichprobe, ausschließlich auf Wirtschaftswaldflächen in Ertrag, angelegt. Die Probepunkte sind in einem Raster von 400x400m verteilt, welches in bestimmten Fällen zur Steigerung der Genauigkeit auf 400x200 oder 200x200m verdichtet werden kann.')
                    obf_doc.doc.add_paragraph('')
                    obf_doc.doc.add_paragraph('Es werden folgende Verfahren zur Erhebung der jeweiligen Daten angewandt:')
                    obf_doc.doc.add_paragraph('Winkelzählprobe (Vorrat)', style='List Bullet')
                    obf_doc.doc.add_paragraph('Fixer Probekreis (Vorrat, Verjüngung)', style='List Bullet')
                    obf_doc.doc.add_paragraph('Schätzung (Bodenvegetation, Verjüngung)', style='List Bullet')
                    obf_doc.doc.add_paragraph('Linienstichprobe (liegendes Totholz)', style='List Bullet')
                    obf_doc.doc.add_paragraph('')
                    obf_doc.doc.add_paragraph('Während der Außenerhebung der Stichprobenteams (Studenten) erfolgt die Einstufung in Schutz- und Wirtschaftswald, sowie in die Umtriebsgruppen auf Basis des Vorgängeroperats. Die Auswertung der erhobenen Daten erfolgt hingegen unter Berücksichtigung der von der Forsteinrichtung neu ausgeschiedenen Waldorte. Sowohl die Einstufung Wirtschaftswald-Schutzwald als auch die Einstufung in die Umtriebsgruppen während der Außenerhebung der Stichproben-Teams (Studenten) wurde basierend auf der alten Forstkarte durchgeführt. Die Auswertung der Ergebnisse wurde mit Überarbeitung der Probeflächendetailliste mit den neuen Waldflächen nach der Taxation durchgeführt.')
                    obf_doc.doc.add_paragraph('')
                    obf_doc.doc.add_paragraph('Dabei werden die Altersklassenflächen jedoch nicht direkt aus der Taxation übernommen, sondern ergeben sich aus dem Altersanteil der einzelnen Baumindividuen in der Winkelzählprobe. Die Ergebnisse der SPI können wegen der unterschiedlichen Altersklassenverteilung erheblich von den Taxationsdaten abweichen.')
                    obf_doc.doc.add_paragraph('')
                    obf_doc.doc.add_paragraph('Die Gesamtergebnisse bzgl. Fläche, Vorrat und Hektarvorrat nach Altersklassen sind hier vergleichend dargestellt.')

                    obf_doc.doc.add_page_break()

                    # Vergleich SPI und TAX
                    obf_doc.doc.add_heading('Vergleich Taxation und Stichprobeninventur',3)

                    # get a list with all FR
                    allfr = obf_fuc.loop_fr()[:,0]

                    # data from TAX
                    table_tax = obf_fuc.fuc_tbl_tax()

                    # data from SPI
                    table_spi = obf_spi.fuc_tbl_spi(obf_fuc.dic.fb, obf_fuc.dic.fr)

                    # loop through all FR
                    for i in np.arange(len(table_tax)):

                        if i == '0':
                            fr = 'All'
                            cod = np.array(['0','0','W'])
                        else:
                            fr = str(allfr[i])
                            cod = np.array([fr,'0','W'])

                        obf_doc.doc.add_heading(obf_fuc.name_code(cod)[1], 4)

                        obf_fuc.fuc_tbl_vergleich(table_tax[i],table_spi[i],fr)
                        # add figure
                        obf_doc.doc.add_picture('tempx.png')
                        obf_doc.docx_paragraph_figure('Vergleich des Vorrates pro Hektar im Wirtschaftswald aus Stichprobe und Taxation im ' + obf_fuc.name_code(cod)[0])

                        table = table_tax[i]

                        obf_doc.docx_paragraph_table('Vorrat, Fläche und Vorrat pro Hektar aus der Taxation, ' + obf_fuc.name_code(cod)[1] + ', Wirtschaftswald, nach Altersklassen (Blöße und AKL I sind unter AKL I zusammengefasst)')
                        obf_doc.docx_table_x(table, Cm(2.2), Cm(2.2), header_rep = True, header = 'Taxationsdaten ' + obf_fuc.name_code(cod)[0], font_size = 7, autofit = True)
                        obf_doc.doc.add_paragraph('')

                        table = table_spi[i]

                        obf_doc.docx_paragraph_table('Vorrat, Fläche und Vorrat pro Hektar aus der Stichprobe, ' + obf_fuc.name_code(cod)[1] + ', Wirtschaftswald, nach Altersklassen (Blöße und AKL I sind unter AKL I zusammengefasst)')
                        obf_doc.docx_table_x(table, Cm(2.2), Cm(2.2), header_rep = True, header = 'Stichprobendaten ' + obf_fuc.name_code(cod)[0], font_size = 7, autofit = True)
                        obf_doc.doc.add_paragraph('')

                    obf_doc.doc.add_page_break()
                except:
                    log = log + "5.2 Vergleich der Vorräte im Wirtschaftswald, Taxation und Stichprobeninventur - error occured\n"
                    self.save_log(log)
                else:
                    log = log + "5.2 Vergleich der Vorräte im Wirtschaftswald, Taxation und Stichprobeninventur - successful\n"
                    self.save_log(log)

            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   5.3 Baumartenausstattung
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            try:
                print('Baumartenausstattung')
                obf_doc.doc.add_heading('Baumartenausstattung', 2)

                for i in code:

                    if np.array_equal(np.array(i), code[0]):
                        obf_doc.doc.add_heading('Revierweise', 3)
                    elif (i[0] == '0') & (int(i[1]) == obf_fuc.get_min_uz(code)):
                        obf_doc.doc.add_heading('Umtriebszeit', 3)

                    print(i)
                    table = obf_fuc.fuc_tbl_plt_ba_anteile(i)

                    # add figure
                    obf_doc.doc.add_picture('tempx.png')

                    # add Caption to figure
                    obf_doc.docx_paragraph_figure('Baumartenanteile in den Altersklassen nach Flächenanteilen, ' + obf_fuc.name_code(i)[1] + ', '  + obf_fuc.name_code(i)[3] + ' (' + obf_fuc.dic.sl + obf_fuc.dic.sn + ' AKL VIII+ enthält die 8. Altersklasse und älter sowie die Überhälter)')

                    # add Caption to table
                    obf_doc.docx_paragraph_table('Baumartenanteile in den Altersklassen nach Flächenanteilen, ' + obf_fuc.name_code(i)[1] + ', '  + obf_fuc.name_code(i)[3] + ' (' + obf_fuc.dic.sl + obf_fuc.dic.sn + 'AKL VIII+ enthält die 8. Altersklasse und älter sowie die Überhälter)')

                    # add table
                    obf_doc.docx_table_x(table, Cm(2.2), Cm(2.2), False, 'Baumartenanteile im ' + obf_fuc.name_code(i)[1] + ' ' + obf_fuc.name_code(i)[2] + ' [%]', 8, True)
                    obf_doc.doc.add_paragraph('')

                obf_doc.doc.add_page_break()
            except:
                log = log + "5.3 Baumartenausstattung - error occured\n"
                self.save_log(log)
            else:
                log = log + "5.3 Baumartenausstattung - successful\n"
                self.save_log(log)

            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   5.4 Bonitätsverlauf der Baumarten (Mittlere Bonitäten der Baumarten nach Wüchsigkeitsgruppen)
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            try:
                print('Bonitätsverlauf der Baumarten')

                obf_doc.doc.add_heading('Bonitätsverlauf der Baumarten (Mittlere Bonitäten der Baumarten nach Wüchsigkeitsgruppen)', 2)

                table = obf_fuc.fuc_tbl_plt_bonitaetsverlauf()

                # add figure
                obf_doc.doc.add_picture('tempx.png')
                # add Caption to figure
                obf_doc.docx_paragraph_figure('Mittlere Ertragsklassen der Baumarten nach Wüchsigkeitsgruppen im FB ' + obf_fuc.dic.dic_num_fb[obf_fuc.dic.fb] + ', Teiloperat ' + str(obf_fuc.dic.to) + ' (' + obf_fuc.dic.sl + obf_fuc.dic.sn + ')')
                # add Caption to table
                obf_doc.docx_paragraph_table('Mittlere Ertragsklassen der Baumarten nach Wüchsigkeitsgruppen im FB ' + obf_fuc.dic.dic_num_fb[obf_fuc.dic.fb] + ', Teiloperat ' + str(obf_fuc.dic.to) + ' (' + obf_fuc.dic.sl + obf_fuc.dic.sn + ')')
                # add table
                obf_doc.docx_table_x(table, Cm(2.2), Cm(2.2), header_rep = True, header = 'Mittlere Ertragsklassen im TO ' + str(obf_fuc.dic.to) + ' [Vfm/ha/a]', font_size = 7, autofit = True)
                obf_doc.doc.add_paragraph('')

                #obf_doc.doc.add_page_break()
            except:
                log = log + "5.4 Bonitätsverlauf der Baumarten (Mittlere Bonitäten der Baumarten nach Wüchsigkeitsgruppen) - error occured\n"
                self.save_log(log)
            else:
                log = log + "5.4 Bonitätsverlauf der Baumarten (Mittlere Bonitäten der Baumarten nach Wüchsigkeitsgruppen) - successful\n"
                self.save_log(log)

            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   5.5 Mittelwerte (Alter, Ertragsklasse, Bestockungsgrad)
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            try:
                print('Mittelwerte')
                obf_doc.doc.add_heading('Mittelwerte (Alter, Ertragsklasse, Bestockungsgrad)', 2)

                category_mittel = np.array(('Schichtalter', 'Ertragsklasse', 'BaumartenBestockgrad'))
                code = obf_fuc.loop_be()

                for j in category_mittel:

                    print(j)
                    obf_doc.doc.add_heading(obf_fuc.dic.dic_avg_head[j], 3)

                    for i in code:

                        print(i)

                        table = obf_fuc.fuc_tbl_mittel(j, i)

                        obf_doc.docx_paragraph_table(obf_fuc.dic.dic_avg_head[j] + ' der Baumarten in den Altersklassen im ' + obf_fuc.name_code(i)[3] + ', FB ' + obf_fuc.dic.dic_num_fb[obf_fuc.dic.fb] + ', Teiloperat ' + str(obf_fuc.dic.to) + ' (AKL VIII+ enthält die 8. Altersklasse und älter sowie die Überhälter)')

                        # add table
                        obf_doc.docx_table_x(table, Cm(1.5), Cm(1.5), header_rep = True, header = obf_fuc.dic.dic_avg_head[j] + ' im TO ' + str(obf_fuc.dic.to) + ' ' + obf_fuc.name_code(i)[2] + obf_fuc.dic.dic_avg_unit[j], font_size = 7, autofit = True)
                        obf_doc.doc.add_paragraph('')

                        if i[2] == 'W':
                            obf_doc.doc.add_page_break()

                    obf_doc.doc.add_page_break()
            except:
                log = log + "5.5 Mittelwerte (Alter, Ertragsklasse, Bestockungsgrad) - error occured\n"
                self.save_log(log)
            else:
                log = log + "5.5 Mittelwerte (Alter, Ertragsklasse, Bestockungsgrad) - successful\n"
                self.save_log(log)

            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   5.6 Forstschäden
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###

            if dat_spi == True:
                try:

                    print('Forstschaeden')
                    obf_doc.doc.add_heading('Forstschäden', 2)
                    obf_doc.doc.add_paragraph('')
                    obf_doc.doc.add_paragraph('Die Auswertungen in diesem Kapitel basieren auf Daten der Stichprobeninventur ' + self.fe_year + ' für des Teiloperat ' + str(obf_fuc.dic.to) + '.')

                    obf_doc.doc.add_heading('Wildschäden – Verbiss', 3)
                    obf_doc.doc.add_paragraph('*** xAbbildung ***')
                    obf_doc.docx_paragraph_figure('Leittriebverbiss auf den Probeflächen des JVSM im TO ' + str(obf_fuc.dic.to))

                    obf_doc.doc.add_heading('Wildschäden – Schälschäden', 3)
                    # get table
                    table = obf_spi.fuc_tbl_spi_ss(obf_fuc.dic.to)
                    # add Caption to table
                    obf_doc.docx_paragraph_table('Schälschäden in den BHD-Klassen nach Stammzahl und Vorrat im T ' + str(obf_fuc.dic.to) + '.')
                    # add table
                    # add table
                    x = obf_doc.get_x('schaelschaeden')
                    obf_doc.docx_table_3x(table, x, header_rep = True, header = 'Flächenübersicht [ha]', font_size = 7, autofit = True)

                    #obf_doc.docx_table_x(table, Cm(2), Cm(2), header_rep = True, header = 'Schälschäden', font_size = 7, autofit = True)
                    obf_doc.doc.add_paragraph('')


                    #obf_doc.doc.add_heading('Schaftschäden', 3)
                    for typ in [['Schad.Urs.', 'Schaftschäden'], ['Kronenverl', 'Kronenverlichtung'], ['SchaftQual', 'Schaftqualität']]:

                        obf_doc.doc.add_heading(typ[1], 3)

                        for value in [['Vfm vor Ort', 'Vorratsbezogene'], ['Stz vor Ort', 'Stammzahlbezogene']]:

                            # get table
                            table = obf_spi.fuc_tbl_plt_spi_schaeden(obf_fuc.dic.to, typ=typ[0], title=typ[1], values=value[0])

                            # add figure
                            obf_doc.doc.add_picture('tempx.png')
                            # add Caption to figure
                            obf_doc.docx_paragraph_figure(value[1] + ' ' + typ[1] + ' im FB ' + obf_fuc.dic.dic_num_fb[obf_fuc.dic.fb] + ', Teiloperat ' + str(obf_fuc.dic.to) + ' (AKL VIII+ enthält die 8. Altersklasse und älter sowie die Überhälter)')
                            # add Caption to table
                            obf_doc.docx_paragraph_table(value[1] + ' ' + typ[1] + ' im FB ' + obf_fuc.dic.dic_num_fb[obf_fuc.dic.fb] + ', Teiloperat ' + str(obf_fuc.dic.to) + ' (AKL VIII+ enthält die 8. Altersklasse und älter sowie die Überhälter)')
                            # add table
                            obf_doc.docx_table_x(table, Cm(2.2), Cm(2.2), header_rep = True, header = typ[1] + ' im TO ' + str(obf_fuc.dic.to) + ' [' + value[0] + ']', font_size = 7, autofit = True)
                            obf_doc.doc.add_paragraph('')

                        obf_doc.doc.add_page_break()

                    #obf_doc.doc.add_paragraph('')
                    #obf_doc.doc.add_paragraph('*** xAbbildung ***')
                    #obf_doc.docx_paragraph_figure('Übersichtskarte der Bestände mit frischer Schälung, erhoben im Rahmen der Forsteinrichtung 2016 im FR Mürzsteg')
                    #obf_doc.doc.add_paragraph('')
                    #obf_doc.doc.add_paragraph('*** xAbbildung ***')
                    #obf_doc.docx_paragraph_figure('Übersichtskarte der Schälgrade im FR Mürzsteg, erhoben im Rahmen der Forsteinrichtung 2016')
                    #obf_doc.doc.add_paragraph('')
                    #obf_doc.doc.add_paragraph('*** xAbbildung ***')
                    #obf_doc.docx_paragraph_figure('Neuschälung in Prozent der erhobenen Stammzahl im TO 1311 aus dem unternehmensweiten Schälmonitoring')
                    #obf_doc.doc.add_paragraph('')
                    #obf_doc.docx_paragraph_table('Verteilung der Schälgrade auf die Fläche des Teiloperates')
                    #obf_doc.doc.add_paragraph('*** xTabelle ***')
                    #obf_doc.doc.add_paragraph('')

                    #obf_doc.doc.add_heading('Kronenverlichtung', 3)
                    ### xxx 2016
                    #obf_doc.doc.add_paragraph('Die Auswertungen in diesem Kapitel basieren auf Daten der Stichprobeninventur 2016 für des Teiloperat ' + str(obf_fuc.dic.to) + '.')
                    #obf_doc.docx_paragraph_table('Hektarbezogene Kronenverlichtung im TO ' + str(obf_fuc.dic.to) + '.')
                    #obf_doc.doc.add_paragraph('*** xTabelle ***')

                    #obf_doc.doc.add_page_break()
                except:
                    log = log + "5.6 Forstschäden - error occured\n"
                    self.save_log(log)
                else:
                    log = log + "5.6 Forstschäden - successful\n"
                    self.save_log(log)

            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   5.7 Einteilung der Umtriebsgruppen
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            try:
                print('Einteilung der Umtriebsgruppen')
                obf_doc.doc.add_heading('Einteilung der Umtriebsgruppen', 2)
                obf_doc.doc.add_heading('Umtriebsgruppen nach Wüchsigkeit', 3)

                table = obf_fuc.fuc_tbl_plt_umtriebsgruppen_wuechsigkeit()

                # add figure
                obf_doc.doc.add_picture('tempx.png')
                # add Caption to figure
                obf_doc.docx_paragraph_figure('Verteilung der Wüchsigkeitsgruppen nach Umtriebsgruppen in Prozent, FB ' + obf_fuc.dic.dic_num_fb[obf_fuc.dic.fb] + ', Teiloperat ' + str(obf_fuc.dic.to))
                # add Caption to table
                obf_doc.docx_paragraph_table('Verteilung der Wüchsigkeitsgruppen nach Umtriebsgruppen in Hektar und Prozent, FB ' + obf_fuc.dic.dic_num_fb[obf_fuc.dic.fb] + ', Teiloperat ' + str(obf_fuc.dic.to))
                # add table
                x = obf_doc.get_x('UZ', obf_fuc.dic.uz, obf_fuc.dic.to)
                obf_doc.docx_table_3x(table, x, header_rep = True, header = 'Wüchsigkeit', font_size = 7, autofit = True)
                obf_doc.doc.add_paragraph('')

                obf_doc.doc.add_page_break()

                obf_doc.doc.add_heading('Umtriebsgruppen nach Standortseinheiten', 3)
                table = obf_fuc.fuc_tbl_plt_umtriebsgruppen_stoe()

                # add figure
                obf_doc.doc.add_picture('tempx.png')
                # add Caption to figure
                obf_doc.docx_paragraph_figure('Verteilung der Standortseinheiten in den Umtriebsgruppen in Prozent, im FB ' + obf_fuc.dic.dic_num_fb[obf_fuc.dic.fb] + ', Teiloperat ' + str(obf_fuc.dic.to) + ', die Bedeutung der Standortseinheiten ist im Kapiltel 3.5.1 Standortseinheiten im Teiloperat ersichtilich')
                # add Caption to table
                obf_doc.docx_paragraph_table('Verteilung der Standortseinheiten in den Umtriebsgruppen in Hektar und Prozent, im FB ' + obf_fuc.dic.dic_num_fb[obf_fuc.dic.fb] + ', Teiloperat ' + str(obf_fuc.dic.to) + ', die Bedeutung der Standortseinheiten ist im Kapiltel 3.5.1 Standortseinheiten im Teiloperat ersichtilich')
                # add table
                x = obf_doc.get_x('UZ', obf_fuc.dic.uz, obf_fuc.dic.to)
                obf_doc.docx_table_3x(table, x, header_rep = True, header = 'Standortseinheiten', font_size = 7, autofit = True)
                obf_doc.doc.add_paragraph('')
            except:
                log = log + "5.7 Einteilung der Umtriebsgruppen - error occured\n"
                self.save_log(log)
            else:
                log = log + "5.7 Einteilung der Umtriebsgruppen - successful\n"
                self.save_log(log)


        #######################################################################################################################
        ###   6 Holzernte
        #######################################################################################################################

        if self.do_sections['6_Holzernte'] == 1:

            print('***   6 Holzernte   ***')
            obf_doc.doc.add_heading('Holzernte', 1)

            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   6.1 Beurteilung der abgelaufenen Periode
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###

            print('Beurteilung der abgelaufenen Periode')
            obf_doc.doc.add_heading('Beurteilung der abgelaufenen Periode', 2)

            ###-----------------------------------------------------------------------------------------------------------------###
            ###   6.1.1 Einschlagsübersicht
            ###-----------------------------------------------------------------------------------------------------------------###

            if dat_es == True:
                try:

                    print('   Einschlagsübersicht')
                    obf_doc.doc.add_heading('Einschlagsübersicht', 3)

                    for i in [['EN','Endnutzung'], ['VN','Vornutzung']]:

                        ### GET 'TO' FROM FILE
                        ### ADD GES = EN + VN
                        ### beschriften von Vornutzung und Endnutzung

                        if len(self.to_old) == 1:
                            text_block = ' für das Teiloperat '
                        else:
                            text_block = ' für die Teiloperate '

                        # create table
                        table = obf_es.fuc_tbl_hiebsatzbilanz(self.to_old, obf_fuc.dic.fr, i)

                        # add figure
                        obf_doc.doc.add_picture('tempx.png')

                        # add Caption to figure
                        obf_doc.docx_paragraph_figure('Gegenüberstellung von Einschlag und Hiebssatz in der ' + i[1] + text_block + self.to_old_text + ' in der abgelaufenen Periode im Wirtschafts- und Schutzwald')

                        # add paragraph
                        obf_doc.docx_paragraph_table('Gegenüberstellung von Einschlag und Hiebssatz in der ' + i[1] + text_block + self.to_old_text + ' in der abgelaufenen Periode im Wirtschafts- und Schutzwald')

                        # add table
                        x = obf_doc.get_x('HS', obf_fuc.dic.uz, obf_fuc.dic.to)
                        obf_doc.docx_table_3x(table, x, header_rep = True, header = 'Hiebssatzbilanz - ' + i[1] + ' [Efm]', font_size = 7, autofit = True)
                        obf_doc.doc.add_paragraph('')

                    obf_doc.doc.add_page_break()

                except:
                    log = log + "6.1.1 Einschlagsübersicht - error occured\n"
                    self.save_log(log)
                else:
                    log = log + "6.1.1 Einschlagsübersicht - successful\n"
                    self.save_log(log)

            ###-----------------------------------------------------------------------------------------------------------------###
            ###   6.1.2 Vergleich Einschlag zu Hiebssatz
            ###-----------------------------------------------------------------------------------------------------------------###

            if dat_es_hs == True:
                try:

                    print('   Vergleich Einschlag zu Hiebssatz')
                    obf_doc.doc.add_heading('Vergleich Einschlag zu Hiebssatz', 3)

                    # get unique FR and rename them
                    fr_name = [obf_fuc.dic.dic_num_fr[i] for i in obf_fuc.dic.fr]
                    # set FR in obf_es_hs Class
                    obf_es_hs.set_fr(fr_name)

                    for j in ['Altersgruppe 1. Schicht', 'Neigungsgruppe (%)', 'Seehöhen Gruppe', 'Umtriebszeit']:

                        obf_doc.doc.add_heading(obf_fuc.dic.dic_es_hs_head[j], 4)

                        unit = obf_fuc.dic.dic_es_hs_unit[j]

                        for i in ['Endnutzung', 'Vornutzung']:
                            # create figure
                            obf_es_hs.plot_es_hs(j, i)
                            # add figure
                            obf_doc.doc.add_picture('tempx.png')
                            # add Caption to figure
                            obf_doc.docx_paragraph_figure('Gegenüberstellung von Einschlag und waldbaulichem Hiebssatz in der ' + i + ' nach ' + obf_fuc.dic.dic_es_hs_head[j] + text_block + self.to_old_text + ' in der abgelaufenen Periode im Wirtschafts- und Schutzwald')

                        obf_doc.doc.add_page_break()
                except:
                    log = log + "6.1.2 Vergleich Einschlag zu Hiebssatz - error occured\n"
                    self.save_log(log)
                else:
                    log = log + "6.1.2 Vergleich Einschlag zu Hiebssatz - successful\n"
                    self.save_log(log)

            ###-----------------------------------------------------------------------------------------------------------------###
            ###   6.1.3 Schadholz
            ###-----------------------------------------------------------------------------------------------------------------###

            if dat_zv_ze == True:
                try:

                    print('   Schadholz')
                    obf_doc.doc.add_heading('Schadholz', 3)

                    obf_zv_ze.set_fr(fr_name)

                    ### last two figures Forstrevier

                    for j in ['Nutzungsart' , 'Forstrevier']:

                        if j == 'Nutzungsart':
                            obf_doc.doc.add_heading('Schadholz nach Nutzungsart', 4)
                        else:
                            obf_doc.doc.add_heading('Schadholz nach Revier', 4)

                        for i in ['Endnutzung', 'Vornutzung']:

                            if (i=='Endnutzung') & (j == 'Nutzungsart'):
                                text_block_legende = '; EA = Schäden andere, ED = Schäden Käfer, EF = Schäden Fangbäume, ES = Schäden Schnee, EW = Schäden Wind'
                                #obf_doc.docx_table_leg(obf_text.leg_es_ze, Cm(1), Cm(5))
                            elif (i=='Vornutzung') & (j == 'Nutzungsart'):
                                text_block_legende = '; VA = Schäden andere, VD = Schäden Käfer, VS = Schäden Schnee, VW = Schäden Wind'
                                #obf_doc.docx_table_leg(obf_text.leg_es_zv, Cm(1), Cm(5))
                            else:
                                text_block_legende = ''

                            # create table
                            table = obf_zv_ze.fuc_tbl_schadholz(j, i, obf_fuc.dic.to)

                            # add figure
                            obf_doc.doc.add_picture('tempx.png')

                            # add Caption to figure
                            obf_doc.docx_paragraph_figure('Schadholzmassen in der ' + i + text_block
                                                          + self.to_old_text + ', in der abgelaufenen Periode im Wirtschafts- und Schutzwald' + text_block_legende)

                            # add paragraph
                            obf_doc.docx_paragraph_table('Schadholzmassen in der ' + i + text_block + self.to_old_text + text_block_legende)

                            # add table
                            obf_doc.docx_table_x(table, Cm(2.2), Cm(2.2), True, 'Schadholzmassen in der ' + i + ' [Efm]')
                            obf_doc.doc.add_paragraph('')

                    obf_doc.doc.add_page_break()

                except:
                    log = log + "6.1.3 Schadholz - error occured\n"
                    self.save_log(log)
                else:
                    log = log + "6.1.3 Schadholz - successful\n"
                    self.save_log(log)

            ###-----------------------------------------------------------------------------------------------------------------###
            ###   6.1.4 Waldbauliche Beurteilung
            ###-----------------------------------------------------------------------------------------------------------------###

            print('   Waldbauliche Beurteilung')
            obf_doc.doc.add_heading('Waldbauliche Beurteilung', 3)
            obf_doc.doc.add_paragraph('***   freier Text   ***')
            obf_doc.doc.add_paragraph('')

            log = log + "6.1.4 Waldbauliche Beurteilung - successful\n"
            self.save_log(log)

            obf_doc.doc.add_page_break()


            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   6.2 Hiebssatzermittlung
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###

            print('Hiebssatzermittlung')
            obf_doc.doc.add_heading('Hiebssatzermittlung', 2)

            ###-----------------------------------------------------------------------------------------------------------------###
            ###   6.2.1 Waldbaulicher Hiebssatz
            ###-----------------------------------------------------------------------------------------------------------------###

            try:

                print('   Waldbaulicher Hiebssatz')
                obf_doc.doc.add_heading('Waldbaulicher Hiebssatz', 3)

                #for i in l:
                table = obf_fuc.fuc_tbl_hs_waldbau('Forstrevier')

                # add paragraph
                obf_doc.docx_paragraph_table('Waldbaulicher dezennaler Hiebssatz in den Revieren und im Teiloperat ' + str(obf_fuc.dic.to))

                # add table
                x = obf_doc.get_x('festgesetzterHS')
                obf_doc.docx_table_3x(table, x, header_rep = True, header = 'Waldbauliche Hiebsätze [Efm]', font_size = 7, autofit = True)
                obf_doc.doc.add_paragraph('')

                table = obf_fuc.fuc_tbl_hs_waldbau('Umtriebszeit')

                # add paragraph
                obf_doc.docx_paragraph_table('Waldbaulicher dezennaler Hiebssatz in den Umtriebsgruppen und im Teiloperat ' + str(obf_fuc.dic.to))

                # add table
                x = obf_doc.get_x('festgesetzterHS')
                obf_doc.docx_table_3x(table, x, header_rep = True, header = 'Waldbauliche Hiebsätze [Efm]', font_size = 7, autofit = True)
                obf_doc.doc.add_paragraph('')

                # add section break
                new_section = obf_doc.doc.add_section(WD_SECTION.NEW_PAGE)
                new_section.start_type

                for i in obf_fuc.dic.fr:

                    table = obf_fuc.fuc_tbl_hs_waldbau_bkl(i)
                    obf_doc.docx_paragraph_table('Waldbaulicher dezennaler Hiebssatz in den Betriebsklassen im FR ' + str(i) )
                    # add table
                    x = obf_doc.get_x('festgesetzterHS_BKL')
                    obf_doc.docx_table_3x(table, x, header_rep = True, header = 'Waldbauliche Hiebsätze im FR ' + str(i) + ' [Efm]', font_size = 7, autofit = True)
                    obf_doc.doc.add_paragraph('')

                # add section break
                new_section = obf_doc.doc.add_section(WD_SECTION.NEW_PAGE)
                new_section.start_type

            except:
                log = log + "6.2.1 Waldbaulicher Hiebssatz - error occured\n"
                self.save_log(log)
            else:
                log = log + "6.2.1 Waldbaulicher Hiebssatz - successful\n"
                self.save_log(log)

            ###-----------------------------------------------------------------------------------------------------------------###
            ###   6.2.2 Berechneter Hiebssatz
            ###-----------------------------------------------------------------------------------------------------------------###

            print('   Berechneter Hiebssatz')
            obf_doc.doc.add_heading('Berechneter Hiebssatz', 3)

            obf_doc.doc.add_heading('Rechnerische Hiebssatzermittlung', 4)

            obf_doc.doc.add_paragraph('Die Herleitung des Hiebssatzes erfolgt unter Berücksichtigung des waldbaulichen und des Nachhaltigkeits-Hiebssatzes. Der waldbauliche Hiebssatz ergibt sich aus der Summe der im Zuge der Taxation geplanten Nutzungseingriffe (Vor- und Endnutzung) unter Berücksichtigung der waldbaulichen und ökonomischen Bewirtschaftungsgrundsätze. Um den Aspekt der Nachhaltigkeit bei der Hiebssatzermittlung zu berücksichtigen, wird auch ein Nachhaltigkeits-Hiebssatz ermittelt. Hierzu kommen folgende Hiebssatzweiser zum Einsatz:')

            obf_doc.doc.add_heading('Hiebssatzweiser für die Endnutzung im Wirtschaftswald', 4)

            for i in [120, 140]:
                obf_doc.doc.add_heading(str(i) + ' Jahre', 5)

                obf_doc.docx_paragraph_table('Eingangsgrößen für die Hiebssatzberechnung')
                obf_doc.doc.add_paragraph('*** xTabelle ***')
                obf_doc.doc.add_paragraph('')

                obf_doc.doc.add_paragraph('*** xAbbildung ***')
                obf_doc.docx_paragraph_figure('Verteilung der Altersklassen nach konkreten Flächen, die schwarze Linie zeigt die Verteilung der Altersklassen gemäß dem Normalwaldmodell')
                obf_doc.doc.add_paragraph('')
                obf_doc.docx_paragraph_table('Fläche, Vorrat und HDZ in den Altersklassen (Wirtschaftswald in Ertrag, U ' + str(i) + ') im Teiloperat')
                obf_doc.doc.add_paragraph('*** xTabelle ***')
                obf_doc.doc.add_paragraph('')
                obf_doc.docx_paragraph_table('Berechnete dezennale Endnutzungshiebsätze in Efm (Wirtschaftswald in Ertrag, U ' + str(i) + ') in Abhängigkeit des Ausgleichszeitraums')
                obf_doc.doc.add_paragraph('*** xTabelle ***')
                obf_doc.doc.add_paragraph('')
                obf_doc.doc.add_paragraph('*** xAbbildung ***')
                obf_doc.docx_paragraph_figure('Verlauf der berechneten dezennalen Endnutzungshiebsätze in Efm (Wirtschaftswald in Ertrag, U ' + str(i) + ') in Abhängigkeit des Ausgleichszeitraums')
                obf_doc.doc.add_paragraph('')

            log = log + "6.2.2 Berechneter Hiebssatz - successful\n"
            self.save_log(log)

            ###-----------------------------------------------------------------------------------------------------------------###
            ###   6.2.3 Festgesetzter Hiebssatz
            ###-----------------------------------------------------------------------------------------------------------------###

            print('   Festgesetzter Hiebssatz')
            obf_doc.doc.add_heading('Festgesetzter Hiebssatz', 3)

            if dat_hs == True:
                try:

                    print('      Hiebssätze FR')
                    obf_doc.doc.add_paragraph('Die festgesetzten Vor- und Endnutzungs-Hiebsätze im Schutzwald sowie die Vornutzungs-Hiebsätze im Wirtschaftswald entsprechen weitgehend den waldbaulichen Hiebsätzen aus der Taxation. Der festgesetzte Endnutzungs-Hiebssatz im Wirtschaftswald ergeht aus Taxation und Berechnung. Eine Begründung des festgesetzten Hiebssatzes erfolgt in den kommenden Kapiteln.')

                    #for i in l:
                    table = obf_hs.fuc_tbl_hs_fest('Forstrevier')

                    # add paragraph
                    obf_doc.docx_paragraph_table('Festgesetzter dezennaler Hiebssatz in den Revieren und im Teiloperat ' + str(obf_fuc.dic.to))

                    # add table
                    x = obf_doc.get_x('festgesetzterHS')
                    obf_doc.docx_table_3x(table, x, header_rep = True, header = 'Festgesetzte Hiebssätze [Efm]', font_size = 7, autofit = True)
                    obf_doc.doc.add_paragraph('')

                    print('      Hiebssätze UZ')
                    table = obf_hs.fuc_tbl_hs_fest('Umtriebszeit')

                    # add paragraph
                    obf_doc.docx_paragraph_table('Festgesetzter dezennaler Hiebssatz in den Umtriebsgruppen und im Teiloperat ' + str(obf_fuc.dic.to))

                    # add table
                    x = obf_doc.get_x('festgesetzterHS')
                    obf_doc.docx_table_3x(table, x, header_rep = True, header = 'Festgesetzter Hiebssatz [Efm]', font_size = 7, autofit = True)
                    obf_doc.doc.add_paragraph('')

                    # add section break
                    new_section = obf_doc.doc.add_section(WD_SECTION.NEW_PAGE)
                    new_section.start_type

                    for i in obf_fuc.dic.fr:

                        table = obf_hs.fuc_tbl_hs_fest_bkl(i)
                        obf_doc.docx_paragraph_table('Festgesetzter dezennaler Hiebssatz in den Betriebsklassen im FR ' + str(i) )

                        # add table
                        x = obf_doc.get_x('festgesetzterHS_BKL')
                        obf_doc.docx_table_3x(table, x, header_rep = True, header = 'Festgesetzte Hiebssätze im FR ' + str(i) + ' [Efm]', font_size = 7, autofit = True)
                        obf_doc.doc.add_paragraph('')

                    # add section break
                    new_section = obf_doc.doc.add_section(WD_SECTION.NEW_PAGE)
                    new_section.start_type

                except:
                    log = log + "6.2.3 Festgesetzter Hiebssatz - error occured\n"
                    self.save_log(log)
                else:
                    log = log + "6.2.3 Festgesetzter Hiebssatz - successful\n"
                    self.save_log(log)

            ###-----------------------------------------------------------------------------------------------------------------###
            ###   6.2.4 Hiebssatz Begründung
            ###-----------------------------------------------------------------------------------------------------------------###

            print('   Hiebssatz Begründung')
            obf_doc.doc.add_heading('Hiebssatz Begründung', 3)
            obf_doc.doc.add_paragraph('Hiebssatz Begründung Text')


            ###-----------------------------------------------------------------------------------------------------------------###
            ###   6.2.5 Vergleich Hiebssatz vergangene und aktuelle Periode
            ###-----------------------------------------------------------------------------------------------------------------###

            print('   Vergleich Hiebssatz vergangene und aktuelle Periode')
            obf_doc.doc.add_heading('Vergleich Hiebssatz vergangene und aktuelle Periode', 3)

            obf_doc.docx_paragraph_table('Vergleich des dezennalen Hiebssatzes in der Vorperiode (HS alt) mit der Periode ' + obf_fuc.dic.laufzeit +' (HS neu)')
            obf_doc.doc.add_paragraph('*** xTabelle ***')
            obf_doc.doc.add_paragraph('')

            obf_doc.doc.add_page_break()

            log = log + "6.2.5 Vergleich Hiebssatz vergangene und aktuelle Periode - successful\n"
            self.save_log(log)

            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   6.3 Ziele und Planung für die Periode
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###

            print('Ziele und Planung für die Periode ' + obf_fuc.dic.laufzeit)
            obf_doc.doc.add_heading('Ziele und Planung für die Periode ' + obf_fuc.dic.laufzeit, 2)

            ###-----------------------------------------------------------------------------------------------------------------###
            ###   6.3.1 Waldbauliche Grundsätze
            ###-----------------------------------------------------------------------------------------------------------------###

            print('Waldbauliche Grundsätze')
            obf_doc.doc.add_heading('Waldbauliche Grundsätze', 3)

            obf_text.fuc_txt_waldbau_grundsatz(obf_doc.doc)

            obf_doc.doc.add_page_break()

            log = log + "6.3.1 Waldbauliche Grundsätze - successful\n"
            self.save_log(log)

            ###-----------------------------------------------------------------------------------------------------------------###
            ###   6.3.2 Nutzungsprofil
            ###-----------------------------------------------------------------------------------------------------------------###

            try:

                print('   Nutzungsprofil')
                obf_doc.doc.add_heading('Nutzungsprofil', 3)

                for i in obf_text.fuc_loop_nutz():

                    obf_doc.doc.add_heading('Nutzungsplanung nach ' + i[1], 4)

                    #'Bewpfl.', 'Zeitpunkt', 'Rückungsart', 'Schlägerungsart'

                    #if i[0] != 'TAX: Altersklasse':
                        #obf_doc.docx_table_leg(i[2], Cm(1), Cm(5))

                    for j in [['EN','Endnutzung'],['VN','Vornutzung']]:

                        # create table
                        table = obf_fuc.fuc_tbl_nutzungszeitpunkt(i[0], j[0], 'Nutzungssumme')
                        txt = obf_dic.nutz_legend(table.iloc[:-1,0].unique(),i[1])

                        # add Caption to table
                        obf_doc.docx_paragraph_table(j[1] + 'smengen nach ' + i[1] + ' in Erntefestmetern in den Revieren und im Teiloperat ' + str(obf_fuc.dic.to) + txt)

                        # add table
                        x = obf_doc.get_x('FR', obf_fuc.dic.fr, obf_fuc.dic.to)
                        obf_doc.docx_table_3x(table, x, header_rep = True, header = j[1] + 'smengen nach ' + i[1], font_size = 7, autofit = True)
                        obf_doc.doc.add_paragraph('')


                    if (i[0] == 'Nutzdringlichkeit') or (i[0] == 'Maßnahmenart'):

                        for j in [['EN','Endnutzung'],['VN','Vornutzung']]:

                            # create table
                            table = obf_fuc.fuc_tbl_nutzungszeitpunkt(i[0], j[0], 'Angriffsfläche')

                            # add Caption to table
                            obf_doc.docx_paragraph_table('Angriffsfläche in der ' + j[1] + ' nach ' + i[1] + ' in Hektar in den Revieren und im Teiloperat ' + str(obf_fuc.dic.to))

                            # add table
                            x = obf_doc.get_x('FR', obf_fuc.dic.fr, obf_fuc.dic.to)
                            obf_doc.docx_table_3x(table, x, header_rep = True, header = 'Angriffsfläche in der ' + j[1] + ' nach ' + i[1], font_size = 7, autofit = True)
                            obf_doc.doc.add_paragraph('')

                obf_doc.doc.add_page_break()

            except:
                log = log + "6.3.2 Nutzungsprofil - error occured\n"
                self.save_log(log)
            else:
                log = log + "6.3.2 Nutzungsprofil - successful\n"
                self.save_log(log)


        #######################################################################################################################
        ###   7 Waldpflege
        #######################################################################################################################

        if self.do_sections['7_Waldpflege'] == 1:

            print('***   7 Waldpflege   ***')
            obf_doc.doc.add_heading('Waldpflege', 1)

            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   7.1 Beurteilung der abgelaufenen Periode
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###

            print('Beurteilung der abgelaufenen Periode')
            obf_doc.doc.add_heading('Beurteilung der abgelaufenen Periode', 2)

            ###-----------------------------------------------------------------------------------------------------------------###
            ###   7.1.1 Waldbauliche Beurteilung
            ###-----------------------------------------------------------------------------------------------------------------###

            print('   Waldbauliche Beurteilung')
            obf_doc.doc.add_heading('Waldbauliche Beurteilung', 3)
            obf_doc.doc.add_paragraph('***   freier Text   ***')
            obf_doc.doc.add_paragraph('')

            log = log + "7.1.1 Waldbauliche Beurteilung - successful\n"
            self.save_log(log)

            obf_doc.doc.add_page_break()

            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   7.1.2 Waldpflegeplan
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###

            if dat_wpplan == True:
                try:
                    print('   Waldpflegeplan')

                    # add section break
                    new_section = obf_doc.doc.add_section(WD_SECTION.NEW_PAGE)
                    new_section.start_type

                    obf_doc.doc.add_heading('Waldpflegeplan', 3)

                    categories = obf_fuc.wpplan_loop()

                    for i in categories:

                        print('      ' + i)

                        table = obf_wp.fuc_tbl_waldpflegeplan(i)

                        # add Caption to table
                        obf_doc.docx_paragraph_table('Auszug aus dem SAP-Waldpflegeplan: Soll-Ist-Vergleich für die verfügbare Periode 2009-2013 (Datenbezug erst ab 2009 -2013, also die letzten 5 Jahre abfragbar, im SAP)')

                        # add table
                        obf_doc.docx_table_x(table, Cm(2.2), Cm(2.2), True, 'Waldpflegeplan ' + i)
                        obf_doc.doc.add_paragraph('')

                    # add section break
                    new_section = obf_doc.doc.add_section(WD_SECTION.NEW_PAGE)
                    new_section.start_type

                except:
                    log = log + "7.1 Beurteilung der abgelaufenen Periode - error occured\n"
                    self.save_log(log)
                else:
                    log = log + "7.1 Beurteilung der abgelaufenen Periode - successful\n"
                    self.save_log(log)

            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   7.2 Ziele und Planung für die Periode 2014-2023
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###

            print('Ziele und Planung für die Periode ' + obf_fuc.dic.laufzeit)
            obf_doc.doc.add_heading('Ziele und Planung für die Periode ' + obf_fuc.dic.laufzeit, 2)

            ###-----------------------------------------------------------------------------------------------------------------###
            ###   7.2.1 Bestandesbegründung | 7.2.2 Waldpflege
            ###-----------------------------------------------------------------------------------------------------------------###

            try:
                for i in ['Bestandesbegründung', 'Waldpflege']:

                    print('   ' + i)

                    # add heading
                    obf_doc.doc.add_heading(i, 3)

                    # right filter
                    if i == 'Bestandesbegründung':
                        filterx = ['AF', 'NB', 'EG']
                    elif i == 'Waldpflege':
                        filterx = ['DP', 'JP', 'KE', 'KF', 'PA', 'AS', 'BU']

                    for j in ['Maßnahmenart', 'Nutzdringlichkeit']:

                        # create table
                        table = obf_fuc.fuc_tbl_waldpflege(filterx, j, 'Forstrevier')
                        txt = obf_dic.nutz_legend(table.iloc[:-1,0].unique(),j)

                        # add Caption to table
                        obf_doc.docx_paragraph_table(i + 'maßnahmen nach ' + j + ' in Hektar und in Prozent' + txt)

                        # add table
                        x = obf_doc.get_x('FR', obf_fuc.dic.fr, obf_fuc.dic.to)
                        #x = get_x(data_wp, j)
                        obf_doc.docx_table_3x(table, x, header_rep = True, header = i + ' nach ' + j, font_size = 7, autofit = True)
                        obf_doc.doc.add_paragraph('')

            except:
                log = log + "7.2.1 Bestandesbegründung | 7.2.2 Waldpflege - error occured\n"
                self.save_log(log)
            else:
                log = log + "7.2.1 Bestandesbegründung | 7.2.2 Waldpflege - successful\n"
                self.save_log(log)

            ###-----------------------------------------------------------------------------------------------------------------###
            ###   7.2.3 Erstdurchforstung
            ###-----------------------------------------------------------------------------------------------------------------###

            try:
                print('   Erstdurchforstung')

                obf_doc.doc.add_heading('Erstdurchforstung', 3)

                # create table
                table = obf_fuc.fuc_tbl_waldpflege(['DE'], 'Nutzdringlichkeit', 'Forstrevier')
                txt = obf_dic.nutz_legend(table.iloc[:-1,0].unique(), 'Nutzdringlichkeit')

                # add Caption to table
                obf_doc.docx_paragraph_table(i + 'maßnahmen nach ' + j + ' in Hektar und in Prozent' + txt)

                # add table
                x = obf_doc.get_x('FR', obf_fuc.dic.fr, obf_fuc.dic.to)
                obf_doc.docx_table_3x(table, x, header_rep = True, header = i + ' nach ' + j, font_size = 7, autofit = True)
                obf_doc.doc.add_paragraph('')

            except:
                log = log + "7.2.3 Erstdurchforstung - error occured\n"
                self.save_log(log)
            else:
                log = log + "7.2.3 Erstdurchforstung - successful\n"
                self.save_log(log)

            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   7.3 Baumartenwahl
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###

            print('Baumartenwahl')

            obf_doc.doc.add_heading('Baumartenwahl', 2)
            obf_doc.doc.add_paragraph('*** Text, Tabellen, Grafiken zur Baumartenwahl ***')
            obf_doc.doc.add_paragraph('')

            obf_doc.doc.add_page_break()


        #######################################################################################################################
        ###   8 Einforstungen
        #######################################################################################################################

        if self.do_sections['8_Einforstung'] == 1:

            print('***   8 Einforstung   ***')

            obf_doc.doc.add_heading('Einforstung', 1)
            obf_doc.doc.add_paragraph('*** Text, Tabellen, Grafiken zur Einforstung ***')
            obf_doc.doc.add_paragraph('')

            obf_doc.doc.add_page_break()


        #######################################################################################################################
        ###   9 Wirtschaftsbeschränkungen
        #######################################################################################################################

        print('***   9 Wirtschaftsbeschränkungen   ***')

        obf_doc.doc.add_heading('Wirtschaftsbeschränkungen', 1)

        if (self.do_sections['9_Wirtschaftsbeschraenkungen'] == 1) & obf_fuc.sw:

            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   9.1 Schutzwaldtypen
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###

            try:
                print('Schutzwaldtypen')
                obf_doc.doc.add_heading('Schutzwaldtypen', 2)

                for i in [['0','Schutzwald'],['S','Standortschutzwald'],['O','Objektschutzwald']]:

                    obf_doc.doc.add_heading(i[1], 3)

                    table = obf_fuc.fuc_tbl_schutzwald(i[0])

                    # add Caption to table
                    obf_doc.docx_paragraph_table(i[1] + 'fläche nach Ertragssituation in Hektar in den Revieren und im Teiloperat ' + str(obf_fuc.dic.to))

                    # add table
                    x = obf_doc.get_x('SW', obf_fuc.data_wo_fl[obf_fuc.data_wo_fl['Bewirtschaftungsform']=='S']['Ertragssituation'].unique(), obf_fuc.dic.to)
                    obf_doc.docx_table_3x(table, x, header_rep = True, header = i[1] + 'fläche', font_size = 7, autofit = True)
                    obf_doc.doc.add_paragraph('')

                #if 'SAT' in
                #obf_doc.doc.add_heading('Saatgutbestände', 2)
                obf_doc.doc.add_page_break()

            except:
                log = log + "9.1 Schutzwaldtypen - error occured\n"
                self.save_log(log)
            else:
                log = log + "9.1 Schutzwaldtypen - successful\n"
                self.save_log(log)

            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###
            ###   9.2 Schutzwalderhaltungszustand
            ###+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++###

            try:
                print('Schutzwalderhaltungszustand')

                obf_doc.doc.add_heading('Schutzwalderhaltungszustand', 2)

                ###-----------------------------------------------------------------------------------------------------------------###
                ###   9.2.1 Methodik
                ###-----------------------------------------------------------------------------------------------------------------###

                print('   Methodik')
                obf_doc.doc.add_heading('Methodik', 3)
                obf_doc.doc.add_paragraph('')

                # add standard text
                obf_text.fuc_txt_sez_met_1(obf_doc.doc)
                obf_doc.doc.add_paragraph('')

                # add Caption to table
                obf_doc.docx_paragraph_table('Zahlenschlüssel zur Bewertung von SW Einheiten. * aggregiert über alle Schichten; ** älteste Schicht')

                # add table Zahlenschlüssel
                d = {'Anzahl der Schichten': ['≥ 3', '2', '1'], 'Bewertung': ['+1', '0', '-1']}
                df0 = pd.DataFrame(data=d)
                obf_doc.docx_table_stoe(df0, Cm(3), Cm(3))

                d = {'Bestockungsgrad*': ['> 0,6', '> 0,4 bis ≤ 0,6', '≤ 0,4'], 'Bewertung': ['+1', '0', '-1']}
                df1 = pd.DataFrame(data=d)
                obf_doc.docx_table_stoe(df1, Cm(3), Cm(3))

                d = {'Maximales Schichtalter**': ['≤ 140 Jahren', '> 140 bis ≤ 200', '> 200 Jahre'], 'Bewertung': ['+1', '0', '-1']}
                df2 = pd.DataFrame(data=d)
                obf_doc.docx_table_stoe(df2, Cm(3), Cm(3))

                d = {'Geländeneigung': ['≤ 70 %', '> 70 und ≤ 100 %', '> 100 %'], 'Bewertung': ['+1', '0', '-1']}
                df3 = pd.DataFrame(data=d)
                obf_doc.docx_table_stoe(df3, Cm(3), Cm(3))

                d = {'Schichtanteil bis max. 40 Jahre*': ['> 10 %', '≥ 1 bis ≤ 10 %', '< 1 %'], 'Bewertung': ['+1', '0', '-1']}
                df4 = pd.DataFrame(data=d)
                obf_doc.docx_table_stoe(df4, Cm(3), Cm(3))
                obf_doc.doc.add_paragraph('')

                # add standard text
                obf_text.fuc_txt_sez_met_2(obf_doc.doc)
                obf_doc.doc.add_paragraph('')
                obf_doc.doc.add_page_break()

                ###-----------------------------------------------------------------------------------------------------------------###
                ###   9.2.2 Ergebnisse
                ###-----------------------------------------------------------------------------------------------------------------###

                print('   Ergebnisse')
                obf_doc.doc.add_heading('Ergebnisse', 3)
                obf_doc.doc.add_paragraph('')

                obf_doc.doc.add_heading('Erhaltungszustand', 4)
                obf_doc.doc.add_paragraph('')

                # get table
                table = obf_natur.fuc_tbl_plt_sez_ha()

                # add figure
                obf_doc.doc.add_picture('tempx.png')

                # add Caption to figure
                obf_doc.docx_paragraph_figure('Schutzwalderhaltungszustand in Hektar in den Revieren. Grün = Schutzwirkung für die nächsten 20 Jahre gegeben; Gelb = Schutzwirkung noch gegeben, negative Entwicklungen sind sichtbar; Rot = Unmittelbarer Handlungsbedarf in den nächsten 10 Jahren.')

                # add Caption to table
                obf_doc.docx_paragraph_table('Schutzwalderhaltungszustand in Hektar in den Revieren und im Teiloperat ' + str(obf_fuc.dic.to) + '. Grün = Schutzwirkung für die nächsten 20 Jahre gegeben; Gelb = Schutzwirkung noch gegeben, negative Entwicklungen sind sichtbar; Rot = Unmittelbarer Handlungsbedarf in den nächsten 10 Jahren.')
                # add table
                obf_doc.docx_table_x(table, Cm(3), Cm(3), header_rep = True, header = 'Schutzwalderhaltungszustand [ha]', font_size = 7, autofit = True)
                obf_doc.doc.add_paragraph('')
                obf_doc.doc.add_page_break()

                obf_doc.doc.add_heading('Maßnahmenplanung', 4)
                obf_doc.doc.add_paragraph('')

                for mass in [['WP', 'Waldpflege'], ['VN', 'Vornutzung'], ['EN', 'Endnutzung']]:

                    obf_doc.doc.add_heading(mass[1], 5)
                    obf_doc.doc.add_paragraph('')

                    if mass[0] == 'WP':
                        obf_text.fuc_txt_sez_wp(obf_doc.doc)
                    elif mass[0] == 'VN':
                        obf_text.fuc_txt_sez_vn(obf_doc.doc)
                    elif mass[0] == 'EN':
                        obf_text.fuc_txt_sez_en(obf_doc.doc)

                    obf_doc.doc.add_paragraph('')

                    for fr in obf_fuc.loop_fr():
                        print(fr)

                        table = obf_natur.fuc_tbl_sez_mass_ha(mass[0], int(fr[0]))
                        obf_doc.docx_paragraph_table('Angriffsflächen ' + mass[1] + ' in Hektar, nach Maßnahmenart und Erhaltungszustand ' + obf_fuc.name_code(fr)[1])
                        obf_doc.docx_table_x(table, Cm(3), Cm(3), header_rep = True, header = 'Angriffsflächen nach Schutzwalderhaltungszustand [ha] - ' + obf_fuc.name_code(fr)[0], font_size = 7, autofit = True)
                        obf_doc.doc.add_paragraph('')

                    if mass[0] != 'WP':
                        for fr in obf_fuc.loop_fr():
                            table = obf_natur.fuc_tbl_sez_mass_v(mass[0], int(fr[0]))
                            obf_doc.docx_paragraph_table('Holzanfall ' + mass[1] + ' in Erntefestmeter nach Erhaltungszustand und Rückungsart ' + obf_fuc.name_code(fr)[1])
                            x = obf_doc.get_x('Schutzwalderhaltungszustand')
                            obf_doc.docx_table_3x(table, x, header_rep = True, header = 'Holzanfall nach Schutzwalderhaltungszustand [Efm]', font_size = 7, autofit = True)
                            obf_doc.doc.add_paragraph('')

                    obf_doc.doc.add_page_break()


                #x = obf_doc.get_x('SW', obf_fuc.data_wo_fl[obf_fuc.data_wo_fl['Bewirtschaftungsform']=='S']['Ertragssituation'].unique(), obf_fuc.dic.to)
                #obf_doc.docx_table_3x(table, x, header_rep = True, header = i[1] + 'fläche', font_size = 7, autofit = True)
                #obf_doc.doc.add_paragraph('')

                #if 'SAT' in
                #obf_doc.doc.add_heading('Saatgutbestände', 2)

                    #obf_doc.doc.add_page_break()

            except:
                log = log + "9.2 Schutzwalderhaltungszustand - error occured\n"
                self.save_log(log)
            else:
                log = log + "9.2 Schutzwalderhaltungszustand - successful\n"
                self.save_log(log)



        #######################################################################################################################
        ###   10 Naturschutz
        #######################################################################################################################

        if (dat_natur == True) & (self.do_sections['10_Naturschutz'] == 1):
            try:

                print('***   10 Naturschutz   ***')

                obf_doc.doc.add_heading('Naturschutz', 1)

                # set not filterd data to obf_natur
                obf_natur.set_flaeche(obf_fuc.data)

                obf_doc.doc.add_heading('Allgemeines', 2)
                obf_doc.doc.add_paragraph('')

                # add standard text
                obf_text.fuc_txt_naturschutz(obf_doc.doc)
                obf_doc.doc.add_paragraph('')

                typs = [['NPK', 'Nationalpark'], ['BPK', 'Biosphärenpark'], ['NSG', 'Naturschutzgebiete'], ['LSG', 'Landschaftsschutzgebiete'], ['SOS', 'Sonstige Schutzgebiete'], ['BIO', 'ÖBf-Biotope und Altholzinseln für den Vogelschutz'], ['LRV', 'Lebensraumvernetzung']]

                for typ in typs:

                    if typ[0] in obf_natur.unique_group:

                        print(typ[1])
                        obf_doc.doc.add_heading(typ[1], 2)

                        temp = obf_natur.data[obf_natur.data['AuswGruppe']==typ[0]]
                        kat_typs = temp['AuswKatTyp'].unique()

                        for kat_typ in kat_typs:

                            if kat_typs.size > 1:
                                print('   ' + obf_natur.dic_auswkTyp[kat_typ])
                                obf_doc.doc.add_heading(obf_natur.dic_auswkTyp[kat_typ], 3)

                            # create table
                            table = obf_natur.fuc_tbl_naturschutz(kat_typ)
                            # add Caption to table
                            obf_doc.docx_paragraph_table('Fläche der ' + obf_natur.dic_auswkTyp[kat_typ] + ' in Hektar in den Revieren und im Teiloperat ' + str(obf_fuc.dic.to))
                            # add table
                            x = obf_doc.get_x('natur', extra=obf_natur.dic_auswkTyp[kat_typ])
                            obf_doc.docx_table_3x(table, x, header_rep = True, header = obf_natur.dic_auswkTyp[kat_typ], font_size = 7, autofit = True)

                            #obf_doc.docx_table_x(table, Cm(4), Cm(4), True, obf_natur.dic_auswkTyp[kat_typ], autofit = True)
                            #x = obf_doc.get_x('Natur', cat=obf_natur.dic_auswkTyp[kat_typ])
                            #obf_doc.docx_table_3x(table, x, header_rep = True, header = obf_natur.dic_auswkTyp[kat_typ], font_size = 7, autofit = True)
                            obf_doc.doc.add_paragraph('')

                            if kat_typ == 'BT':
                                # create table
                                table = obf_natur.fuc_tbl_biotop()
                                # add Caption to table
                                obf_doc.docx_paragraph_table('Fläche der Biotoptypen in Hektar in den Revieren und im Teiloperat ' + str(obf_fuc.dic.to))
                                # add table
                                x = obf_doc.get_x('FR', obf_fuc.dic.fr, obf_fuc.dic.to)
                                obf_doc.docx_table_3x(table, x, header_rep = True, header = 'Biotoptypen', font_size = 7, autofit = True)
                                #obf_doc.docx_table_x(table, Cm(4), Cm(3.3), True, 'Biotoptypen', autofit = True)
                                obf_doc.doc.add_paragraph('')

                            elif kat_typ == 'AHI':
                                # create table
                                table = obf_natur.fuc_tbl_birdlife()
                                # add Caption to table
                                obf_doc.docx_paragraph_table('Fläche der einzelnen Altholzinseln in Hektar in den Revieren und im Teiloperat ' + str(obf_fuc.dic.to))
                                # add table
                                obf_doc.docx_table_x(table, Cm(3.5), Cm(3.5), True, 'Altholzinseln', autofit = True)
                                obf_doc.doc.add_paragraph('')

                                obf_doc.doc.add_page_break()

            except:
                log = log + "10.0 Naturschutz - error occured\n"
                self.save_log(log)
            else:
                log = log + "10.0 Naturschutz - successful\n"
                self.save_log(log)


        #######################################################################################################################
        ###   11 Vormerkungen für die nächste Forsteinrichtung
        #######################################################################################################################

        if self.do_sections['11_Vormerkungen'] == 1:

            print('***   11 Vormerkungen für die nächste Forsteinrichtung   ***')

            obf_doc.doc.add_heading('Vormerkungen für die nächste Forsteinrichtung', 1)
            obf_doc.doc.add_paragraph('*** Freier Text ***')

            obf_doc.doc.add_page_break()


        #######################################################################################################################
        ###   12 Anhang
        #######################################################################################################################

        if self.do_sections['12_Anhang'] == 1:

            print('***   12 Anhang   ***')

            obf_doc.doc.add_heading('Anhang', 1)

            obf_doc.doc.add_heading('Verwendete Ertragstafeln', 2)

            for sets in et_sets['Ertragstafelset'].unique():

                # add paragraph
                obf_doc.docx_paragraph_table('Zuordnung der Ertragstafeln zu den Baumarten im Ertragstafelset ' + str(sets) + '' + '')

                table = obf_et.fuc_tbl_et_ba(sets, allba=True)

                # add table
                obf_doc.docx_table_text(table, width=[2,2,5,5,2], header_rep = True, header = 'Ertragstafeln', font_size = 7, autofit = True)
                obf_doc.doc.add_paragraph('')

            obf_doc.doc.add_page_break()


        # save the doc
        path_save = os.path.join(self.data_path, 'TO' + str(obf_fuc.dic.to) + '_operat.docx')
        obf_doc.doc.save(path_save)

        log = log + "12.0 Anhang - successful\n"
        self.save_log(log)

        print('end')
        return('successfully made')
