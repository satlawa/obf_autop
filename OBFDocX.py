################################################################################
#
#   fuc_tbl_hiebsatzbilanz (to, filterx)
#
################################################################################

import numpy as np
import docx
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.shared import Pt
from openpyxl import load_workbook

from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION

class OBFDocX(object):

    def __init__(self, doc):

        self.count_figure = 1
        self.count_table = 1

        self.doc = doc


    ###=========================================================================
    ## get how many 1.000
    ###=========================================================================

    def fuc_tausend(self,c):
        i=0
        while True:
            c = int(c/1000)
            i = i+1
            if c < 1000:
                break
        return i


    ###=========================================================================
    ## format numbers - make . -> , and 1000 -> 1.000
    ###=========================================================================

    def fuc_number_format(self,num):

        string = str(num)
        add = 0

        if '.' in string:
            string=string.replace('.',",")
            add = len(string.split(',',1)[1])+1

        if (num >= 1000) | (num <= -1000):
            for i in range(self.fuc_tausend(num),0,-1):
                x = (-3*i) - add
                string = string[:x] + '.' + string[x:]

        return string

    def Figure(self,paragraph, text):
        run = paragraph.add_run()
        r = run._r
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        r.append(fldChar)
        instrText = OxmlElement('w:instrText')
        instrText.text = ' SEQ Figure \\* ARABIC'
        r.append(instrText)
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        r.append(fldChar)
        run = paragraph.add_run(text)

    def Table(self,paragraph, text):
        run = paragraph.add_run()
        r = run._r
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        r.append(fldChar)
        instrText = OxmlElement('w:instrText')
        instrText.text = ' SEQ Table \\* ARABIC'
        r.append(instrText)
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        r.append(fldChar)
        run = paragraph.add_run(text)


    ###=========================================================================
    ###   docx figure paragraph
    ###=========================================================================

    def docx_paragraph_figure(self, text):

        paragraph = self.doc.add_paragraph('Abbildung ', style='Caption')
        self.Figure(paragraph, ': ' + text)
        self.count_figure += 1

    ###=========================================================================
    ###   docx table paragraph
    ###=========================================================================


    def docx_paragraph_table(self, text):

        paragraph = self.doc.add_paragraph('Tabelle ', style='Caption')
        self.Table(paragraph, ': ' + text)
        self.count_table += 1


    ###=========================================================================
    ## docx table - convert a pandas table to a docx table (docx)
    ###=========================================================================

    def docx_table_text(self, table, width, header_rep = False, header = '', font_size = 7, autofit = True):

        if header_rep == True:
            add_row = 2
        else:
            add_row = 1

        # add table
        t = self.doc.add_table(table.shape[0]+add_row, table.shape[1])

        # set style
        t.style = 'ÖBF_text' #'Light List'

        t.autofit = autofit

        if header_rep == True:
            # add header
            for j in range(table.shape[-1]):
                t.cell(0,j).width = Cm(width[j])

            t.cell(0,0).text = header
            a = t.cell(0, 0)
            b = t.cell(0, table.shape[-1]-1)
            A = a.merge(b)

        # add the header rows.
        for j in range(table.shape[-1]):
            t.cell(add_row-1,j).text = str(table.columns[j])
            t.cell(add_row-1,j).width = Cm(width[j])


        # add the rest of the data frame
        for i in range(table.shape[0]):
            for j in range(table.shape[-1]):

                # change . to ,
                if isinstance(table.values[i,j], str):
                    string = table.values[i,j]
                else:
                    #string = self.fuc_number_format(table.values[i,j])
                    string = str(table.values[i,j])

                t.cell(i+add_row,j).text = string
                t.cell(i+add_row,j).width = Cm(width[j])

        # repeat table header
        if header_rep == True:

            row = t.rows[1]
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            tblHeader = OxmlElement('w:tblHeader')
            tblHeader.set(qn('w:val'), "true")
            trPr.append(tblHeader)

        # set row size and font size
        for row in t.rows:

            # set hight
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(270))
            trHeight.set(qn('w:hRule'),'atLeast')
            trPr.append(trHeight)

            # set font size
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(font_size)

        # make frst header row bigger
        if header_rep == True:
            row = t.rows[0]

            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(300))
            trHeight.set(qn('w:hRule'),'atLeast')
            trPr.append(trHeight)

            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(10)

    ## docx - convert a pandas table to a docx table

    def docx_table_x(self, table, width1, width2, header_rep = False, header = '', font_size = 7, autofit = True):

        if header_rep == True:
            add_row = 2
        else:
            add_row = 1

        # add table
        t = self.doc.add_table(table.shape[0]+add_row, table.shape[1])

        # set style
        t.style = 'ÖBF' #'Light List'

        t.autofit = autofit

        if header_rep == True:
            # add header
            for j in range(table.shape[-1]):
                if j == 0:
                    t.cell(0,j).width = width1
                else:
                    t.cell(0,j).width = width2

            t.cell(0,0).text = header
            a = t.cell(0, 0)
            b = t.cell(0, table.shape[-1]-1)
            A = a.merge(b)

        # add the header rows.
        for j in range(table.shape[-1]):
            t.cell(add_row-1,j).text = str(table.columns[j])
            if j == 0:
                t.cell(add_row-1,j).width = width1
            else:
                t.cell(add_row-1,j).width = width2

        # add the rest of the data frame
        for i in range(table.shape[0]):
            for j in range(table.shape[-1]):

                # change . to ,
                if isinstance(table.values[i,j], str):
                    string = table.values[i,j]
                else:
                    string = self.fuc_number_format(table.values[i,j])

                t.cell(i+add_row,j).text = string
                if j == 0:
                    t.cell(i+add_row,j).width = width1
                else:
                    t.cell(i+add_row,j).width = width2

        # repeat table header
        if header_rep == True:

            row = t.rows[1]
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            tblHeader = OxmlElement('w:tblHeader')
            tblHeader.set(qn('w:val'), "true")
            trPr.append(tblHeader)

        # set row size and font size
        for row in t.rows:

            # set hight
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(270))
            trHeight.set(qn('w:hRule'),'atLeast')
            trPr.append(trHeight)

            # set font size
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(font_size)

        # make frst header row bigger
        if header_rep == True:
            row = t.rows[0]

            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(300))
            trHeight.set(qn('w:hRule'),'atLeast')
            trPr.append(trHeight)

            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(10)

    ## docx - convert a pandas table to a docx table

    def docx_table_3x(self , table, x, header_rep = True, header = '', font_size = 7, autofit = True):

        add_row = 3

        # add table
        t = self.doc.add_table(table.shape[0]+add_row, table.shape[1])

        # set style
        t.style = 'ÖBF_gitter' #'Light List'

        t.autofit = autofit

        # add the 1.header row - Arten.
        for j in range(table.shape[-1]):
            t.cell(0,j).width = Cm(x[1][j])

            t.cell(0,0).text = header
            a = t.cell(0, 0)
            b = t.cell(0, table.shape[-1]-1)
            A = a.merge(b)

        # add the 2.header row - Arten.
        for j in range(table.shape[-1]):
            t.cell(add_row-2,j).width = Cm(x[1][j])

        for j in x[0]:
            if j[0] != 0:
                t.cell(add_row-2, j[0]).text = j[2]
                a = t.cell(add_row-2, j[0])
                b = t.cell(add_row-2, j[1])
                A = a.merge(b)

        # add the 3.header row - Einheiten.
        for j in range(table.shape[-1]):
            t.cell(add_row-1,j).text = str(table.columns[j])
            t.cell(add_row-1,j).width = Cm(x[1][j])

        # add the rest of the data frame
        for i in range(table.shape[0]):
            for j in range(table.shape[-1]):

                # change . to ,
                if isinstance(table.values[i,j], str):
                    string = table.values[i,j]
                else:
                    if j > 0:
                        string = self.fuc_number_format(table.values[i,j])
                    else:
                        if isinstance(table.values[i,j], int):
                            string = str(table.values[i,j])
                        else:
                            string = str(int(table.values[i,j]))

                t.cell(i+add_row,j).text = string
                t.cell(i+add_row,j).width = Cm(x[1][j])

        # repeat table header
        if header_rep == True:

            row = t.rows[1]
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            tblHeader = OxmlElement('w:tblHeader')
            tblHeader.set(qn('w:val'), "true")
            trPr.append(tblHeader)

            row = t.rows[2]
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            tblHeader = OxmlElement('w:tblHeader')
            tblHeader.set(qn('w:val'), "true")
            trPr.append(tblHeader)

        # set row size and font size
        for row in t.rows:

            # set hight
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(270))
            trHeight.set(qn('w:hRule'),'atLeast')
            trPr.append(trHeight)

            # set font size
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(font_size)

        # make frst header row bigger
        if header_rep == True:
            row = t.rows[0]

            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(300))
            trHeight.set(qn('w:hRule'),'atLeast')
            trPr.append(trHeight)

            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(10)

    # esencial

    def get_x(self, kind, cat=0, to=0, extra='Teil'):

        dic_SW = {"I":"in Ertrag", "A":"außer Ertrag"}

        x = [[0]]
        y = [2]

        if ((kind == 'FR') | (kind == 'Forstrevier')):
            j = 1
            for i in cat:
                x = x + [[j,j+1,'FR ' + str(i)]]
                j = j + 2
                y = y + [3, 1.5]
            x = x + [[j,j+1,'TO ' + str(to)]]
            y = y + [3, 1.5]

        if kind == 'UZ':
            j = 1
            for i in cat:
                x = x + [[j,j+1,'U ' + str(i)]]
                j = j + 2
                y = y + [3, 1.5]
            x = x + [[j,j+1,'TO ' + str(to)]]
            y = y + [3, 1.5]

        if kind == 'Nutzdringlichkeit':
            j = 1
            for i in cat:
                x = x + [[j,j+1,'Dringlichkeit ' + str(i)]]
                j = j + 2
                y = y + [3, 1.5]
            x = x + [[j,j+1,'TO ' + str(to)]]
            y = y + [3, 1.5]

        if kind == 'SW':
            j = 1
            for i in cat:
                x = x + [[j,j+1,dic_SW[i]]]
                j = j + 2
                y = y + [3, 1.5]
            x = x + [[j,j+1,'Ges.']]
            y = y + [3, 1.5]

        if kind == 'Natur':
            x = [[0],[1,1,cat],[2,2,'Anteil an Gesamtfläche'], [3,3,'Gesamtfläche']]
            y = [3.8,3.8,3.8,3.8]

        if kind == 'HS':
            x = [[0],[1,3,'Hiebsatz'],[4,6,'Einschlag'],[7,9,'Bilanz']]
            y = [1.2,1.8,1.8,1.8,1.8,1.8,1.8,1.8,1.8,1.8]

        if kind == 'festgesetzterHS':
            x = [[0],[1,1,''],[2,4,'Endnutzung'],[5,7,'Vornutzung'],[8,10,'Gesamt']]
            y = [1.2,1.8,1.8,1.8,1.8,1.8,1.8,1.8,1.8,1.8,1.8]

        if kind == 'festgesetzterHS_BKL':
            x = [[0],[0,0,''],[1,1,''],[2,2,''],[3,3,''],[4,6,'Endnutzung'],[7,9,'Vornutzung'],[10,12,'Gesamt']]
            y = [1.2,1.8,1.8,1.8,1.8,1.8,1.8,1.8,1.8,1.8,1.8,1.8,1.8]

        if kind == 'flaechenuebersicht':
            x = [[0],[1,1,''],[2,6,'Waldfläche'],[7,9,'Nebengründe'],[10,10,'Gesamt-fläche']]
            y = [3,3,2,2,2,2,2,2,2,2,2]

        if kind == 'sw_flaechenuebersicht':
            x = [[0],[1,3,'Standortschutzwald'],[4,6,'Objektschutzwald'],[7,9,'Bannwald'],[10,12,'Schutzwald Gesamt']]
            y = [3,2,2,2,2,2,2,2,2,2,2,2,2]
            #y = [3,2.5,2.5,2.5,2.5,2.5,2.5,2.5,2.5,2.5,,2.5,2.5,2.5]

        if kind == 'schaelschaeden':
            x = [[0],[1,1,'Gesamt'],[2,2,'Schälschäden'],[3,3,'Anteil'],[4,4,'Gesamt'],[5,5,'Schälschäden'],[6,6,'Anteil']]
            y = [3,2,2,2,2,2,2]

        if kind == 'Schutzwalderhaltungszustand':
            x = [[0],[1,3,'Grün'],[4,6,'Gelb'],[7,9,'Rot'],[10,12,'Gesamt']]
            y = [1.2,1.5,1.5,1.5,1.5,1.5,1.5,1.5,1.5,1.5,1.5,1.5,1.5]

        if kind == 'natur':
            x = [[0],[1,1,extra],[2,2,'Anteil'],[3,3,'Gesamt']]
            y = [4,4,4,4]

        z = [x,y]

        return z

    ## docx - convert a pandas table to a docx table

    def docx_table_stoe(self, table, width1, width2):

        header_rep = True

        t = self.doc.add_table(table.shape[0]+1, table.shape[1])

        t.style = 'ÖBF' #'Light List'

        #t.autofit = False

        # add the header rows.
        for j in range(table.shape[-1]):
            t.cell(0,j).text = str(table.columns[j])
            if j == 0:
                t.cell(0,j).width = width1
            else:
                t.cell(0,j).width = width2

        # add the rest of the data frame
        for i in range(table.shape[0]):
            for j in range(table.shape[-1]):

                # change . to ,
                if isinstance(table.values[i,j], str):
                    string = table.values[i,j]
                else:
                    string = self.fuc_number_format(table.values[i,j])

                t.cell(i+1,j).text = string
                if j == 0:
                    t.cell(i+1,j).width = width1
                else:
                    t.cell(i+1,j).width = width2

        if header_rep == True:

            # repeat table header
            row = t.rows[0]
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            tblHeader = OxmlElement('w:tblHeader')
            tblHeader.set(qn('w:val'), "true")
            trPr.append(tblHeader)

        for row in t.rows:

            # set hight
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(270))
            trHeight.set(qn('w:hRule'),'atLeast')
            trPr.append(trHeight)

            # set font size
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(8)
        #if header_rep == True:

            # add header row
            #t_temp = docx.add_table(1, 1)
            #t_temp.style = 'ÖBF'#'ÖBF'
            #t_temp.cell(0,0).text = header
            #t_temp.cell(0,0).width = width1 + width2 * (table.shape[-1]-1)

    def docx_table_leg(self, table, width1, width2):

        t = self.doc.add_table(table.shape[0]+1, table.shape[1])

        t.style = 'ninia'

        t.autofit = False

        # add the header rows.
        for j in range(table.shape[-1]):
            t.cell(0,j).text = str(table.columns[j])
            if j == 0:
                t.cell(0,j).width = width1
            else:
                t.cell(0,j).width = width2

        # add the rest of the data frame
        for i in range(table.shape[0]):
            for j in range(table.shape[-1]):

                # change . to ,
                if isinstance(table.values[i,j], str):
                    string = table.values[i,j]
                else:
                    string = str(table.values[i,j])

                t.cell(i+1,j).text = string
                if j == 0:
                    t.cell(i,j).width = width1
                else:
                    t.cell(i,j).width = width2

        for row in t.rows:

            # set hight
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(270))
            trHeight.set(qn('w:hRule'),'atLeast')
            trPr.append(trHeight)

            # set font size
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(8)

        # make frst header row bigger
        row = t.rows[0]
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(300))
        trHeight.set(qn('w:hRule'),'atLeast')
        trPr.append(trHeight)

        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(10)
